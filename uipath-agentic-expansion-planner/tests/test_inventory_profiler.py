import importlib.util
import json
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
SCRIPT = ROOT / "scripts" / "inventory_profiler.py"
RENDER_SCRIPT = ROOT / "scripts" / "render_executive_docx.py"
VERIFY_SCRIPT = ROOT / "scripts" / "verify_executive_docx.py"
FIXTURE_CSV = ROOT / "tests" / "fixtures" / "inventory.csv"
EXPECTED_PROFILE = ROOT / "tests" / "fixtures" / "expected_profile_summary.json"


def load_module():
    spec = importlib.util.spec_from_file_location("inventory_profiler", SCRIPT)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


def has_python_docx() -> bool:
    return importlib.util.find_spec("docx") is not None


def has_openpyxl() -> bool:
    return importlib.util.find_spec("openpyxl") is not None


class InventoryProfilerTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.module = load_module()

    def write_inventory(self, path: Path) -> None:
        path.write_text(FIXTURE_CSV.read_text(encoding="utf-8"), encoding="utf-8")

    def test_parse_number_supports_planning_shorthand(self):
        cases = {
            "~24k calls": 24_000.0,
            "1.25 million records": 1_250_000.0,
            "2B transactions": 2_000_000_000.0,
            "(3.5 thousand)": -3_500.0,
            "18.5%": 18.5,
        }
        for raw, expected in cases.items():
            with self.subTest(raw=raw):
                self.assertEqual(self.module.parse_number(raw), expected)

    def test_average_handling_minutes_maps_to_handling_time(self):
        guesses = self.module.guess_columns(
            ["Use Case", "Average Handling Minutes", "Annual Transactions"]
        )
        self.assertEqual(guesses["handling_time"]["best"], "Average Handling Minutes")

    def test_negated_production_statuses_do_not_count_as_deployed(self):
        for value in ("Not deployed", "Never live", "Pre-production"):
            with self.subTest(value=value):
                self.assertEqual(self.module.classify_status(value), "pipeline")
                self.assertEqual(
                    self.module.classify_lifecycle_status(value), "pipeline"
                )

    def test_customer_lifecycle_terms_are_consistent_with_recommendation_safety(self):
        cases = {
            "Pipeline": ("pipeline", "pipeline"),
            "Not Approved": ("excluded", "rejected"),
            "Declined": ("excluded", "rejected"),
            "Paused": ("excluded", "paused"),
        }
        for raw, expected in cases.items():
            with self.subTest(raw=raw):
                self.assertEqual(
                    (
                        self.module.classify_status(raw),
                        self.module.classify_lifecycle_status(raw),
                    ),
                    expected,
                )

    def test_safe_source_name_removes_path_and_control_characters(self):
        source = Path("/private/customer/portfolio\nreview.xlsx")
        self.assertEqual(self.module.safe_source_name(source), "portfolio review.xlsx")
        self.assertEqual(
            self.module.safe_source_name(Path("portfolio|review?.xlsx")),
            "portfolio-review-.xlsx",
        )

    def test_source_dates_are_normalized_and_summarized(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / "dated-inventory.csv"
            source.write_text(
                "Name,Description,Status,Department,Last Updated\n"
                "One,First process,Deployed,Finance,2026-06-01\n"
                "Two,Second process,Pipeline,Finance,07/02/2026\n"
                "Three,Third process,Idea,Operations,not-a-date\n"
                "Four,Fourth process,Paused,Operations,\n",
                encoding="utf-8",
            )
            result = subprocess.run(
                [
                    sys.executable,
                    str(SCRIPT),
                    "--input",
                    str(source),
                    "--outdir",
                    str(root / "profile"),
                ],
                capture_output=True,
                text=True,
                check=False,
            )
            self.assertEqual(result.returncode, 0, result.stderr)
            profile = json.loads(
                (root / "profile" / "inventory_profile.json").read_text(
                    encoding="utf-8"
                )
            )
            self.assertEqual(
                profile["metadata"]["source_date_summary"],
                {
                    "columns": ["Last Updated"],
                    "nonblank_count": 3,
                    "valid_count": 2,
                    "invalid_count": 1,
                    "earliest_date": "2026-06-01",
                    "latest_date": "2026-07-02",
                },
            )
            self.assertEqual(
                [item["source_date"] for item in profile["inventory_items"]],
                ["2026-06-01", "2026-07-02", None, None],
            )

    def test_csv_profile_outputs_json_and_markdown_contract(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source = tmp_path / "inventory.csv"
            outdir = tmp_path / "profile"
            self.write_inventory(source)

            result = subprocess.run(
                [
                    sys.executable,
                    str(SCRIPT),
                    "--input",
                    str(source),
                    "--outdir",
                    str(outdir),
                ],
                capture_output=True,
                text=True,
                check=False,
            )

            self.assertEqual(result.returncode, 0, result.stderr)
            profile_path = outdir / "inventory_profile.json"
            markdown_path = outdir / "inventory_profile.md"
            self.assertTrue(profile_path.exists())
            self.assertTrue(markdown_path.exists())

            profile = json.loads(profile_path.read_text(encoding="utf-8"))
            expected = json.loads(EXPECTED_PROFILE.read_text(encoding="utf-8"))
            self.assertEqual(profile["metadata"]["source_file"], "inventory.csv")
            self.assertNotIn(str(tmp_path), json.dumps(profile))
            self.assertEqual(profile["metadata"]["total_rows"], expected["total_rows"])
            self.assertEqual(profile["metadata"]["sheet_count"], expected["sheet_count"])
            self.assertEqual(
                profile["core_field_mapping"]["use_case_name"],
                expected["use_case_name_field"],
            )
            self.assertEqual(profile["core_field_mapping"]["annual_volume"], expected["annual_volume_field"])
            self.assertEqual(
                profile["status_summary"]["normalized_status_counts"],
                expected["status_counts"],
            )
            self.assertEqual(
                profile["data_quality"]["duplicate_name_groups"][0]["normalized_name"],
                expected["duplicate_name"],
            )
            self.assertEqual(profile["numeric_profiles"]["Annual Volume"]["sum"], expected["annual_volume_sum"])
            self.assertEqual(profile["representative_rows"][0]["Use Case Name"], expected["first_representative_name"])

            markdown = markdown_path.read_text(encoding="utf-8")
            self.assertIn("# Inventory profile", markdown)
            self.assertIn("## Detected core field mapping", markdown)
            self.assertIn("| production | 1 |", markdown)
            self.assertIn("Invoice Intake", markdown)

    def test_xlsx_inputs_explain_missing_optional_dependency(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source = tmp_path / "inventory.xlsx"
            source.write_text("not a real workbook", encoding="utf-8")
            result = subprocess.run(
                [
                    sys.executable,
                    str(SCRIPT),
                    "--input",
                    str(source),
                    "--outdir",
                    str(tmp_path / "profile"),
                ],
                capture_output=True,
                text=True,
                check=False,
            )

            if self.module.load_workbook is None:
                self.assertEqual(result.returncode, 1)
                self.assertIn("openpyxl is required", result.stderr)
            else:
                self.assertEqual(result.returncode, 1)

    @unittest.skipUnless(has_openpyxl(), "openpyxl is not installed")
    def test_multi_sheet_workbook_profile_reports_each_sheet(self):
        from openpyxl import Workbook

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / "portfolio.xlsx"
            workbook = Workbook()
            finance = workbook.active
            finance.title = "Finance"
            finance.append(["Name", "Description", "Status", "Department"])
            finance.append(["Invoice intake", "Capture invoice records", "Deployed", "Finance"])
            operations = workbook.create_sheet("Operations")
            operations.append(
                ["Automation Name", "Process Description", "Lifecycle", "Business Unit"]
            )
            operations.append(["Case triage", "Route incoming cases", "Pipeline", "Operations"])
            operations.append([None, None, None, None])
            operations.append(["Archive cleanup", "Retire old records", "Retired", "Operations"])
            workbook.save(source)

            result = subprocess.run(
                [
                    sys.executable,
                    str(SCRIPT),
                    "--input",
                    str(source),
                    "--outdir",
                    str(root / "profile"),
                ],
                capture_output=True,
                text=True,
                check=False,
            )
            self.assertEqual(result.returncode, 0, result.stderr)
            profile = json.loads(
                (root / "profile" / "inventory_profile.json").read_text(encoding="utf-8")
            )
            self.assertEqual(profile["metadata"]["sheet_count"], 2)
            self.assertEqual(profile["metadata"]["total_rows"], 3)
            self.assertEqual(
                [(item["sheet"], item["rows"]) for item in profile["sheets"]],
                [("Finance", 1), ("Operations", 2)],
            )
            items = {item["name"]: item for item in profile["inventory_items"]}
            self.assertEqual(items["Case triage"]["status"], "pipeline")
            self.assertEqual(items["Case triage"]["department"], "Operations")
            self.assertEqual(items["Archive cleanup"]["lifecycle_status"], "retired")
            self.assertEqual(items["Archive cleanup"]["row_number"], 4)
            self.assertTrue(items["Archive cleanup"]["inventory_id"].endswith("R00004"))
            self.assertEqual(
                profile["sheet_field_mappings"]["Operations"]["use_case_name"],
                "Automation Name",
            )
            self.assertEqual(
                profile["status_summary"]["normalized_status_counts"],
                {"production": 1, "pipeline": 1, "excluded": 1},
            )

    def test_profile_rejects_inventory_id_collisions(self):
        row = {
            "Use Case Name": "Example",
            "Description": "Synthetic collision row",
            "Status": "Idea",
            "Department": "Operations",
            "__row_number": 2,
        }
        sheets = {
            "Sheet A": [{**row, "__sheet": "Sheet A"}],
            "Sheet-A": [{**row, "__sheet": "Sheet-A"}],
        }
        with self.assertRaisesRegex(ValueError, "inventory ID collision"):
            self.module.build_profile(Path("inventory.xlsx"), sheets)

    @unittest.skipUnless(has_python_docx(), "python-docx is not installed")
    def test_render_and_verify_docx_when_python_docx_is_available(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            markdown = tmp_path / "brief.md"
            docx = tmp_path / "brief.docx"
            markdown.write_text(
                "## Executive Summary\n\n"
                "Summary.\n\n"
                "## Source and Assumption Note\n\n"
                "Source note.\n\n"
                "## Current Automation Footprint\n\n"
                "Footprint.\n\n"
                "| Area | Signal |\n"
                "| --- | --- |\n"
                "| Finance | Production use cases |\n\n"
                "## Public Strategy Alignment\n\n"
                "Alignment.\n\n"
                "## Prioritized Portfolio\n\n"
                "| Rank | Opportunity | Why now |\n"
                "| --- | --- | --- |\n"
                "| 1 | Invoice Intake | High volume |\n\n"
                "## Top 5 High-Impact Recommendations\n\n"
                "### Recommendation 1: Invoice Intake\n\n"
                "Summary.\n\n"
                "### Recommendation 2: Permit Review\n\n"
                "Summary.\n\n"
                "### Recommendation 3: Exception Triage\n\n"
                "Summary.\n\n"
                "### Recommendation 4: Case Routing\n\n"
                "Summary.\n\n"
                "### Recommendation 5: Audit Prep\n\n"
                "Summary.\n\n"
                "## Top 3 Low-Friction POC Candidates\n\n"
                "### POC 1: Invoice Intake\n\n"
                "Plan.\n\n"
                "### POC 2: Permit Review\n\n"
                "Plan.\n\n"
                "### POC 3: Exception Triage\n\n"
                "Plan.\n\n"
                "## Value Framing\n\n"
                "| Metric | Value |\n"
                "| --- | --- |\n"
                "| Annual volume | 42,000 |\n\n"
                "## Deployment and Governance Considerations\n\n"
                "Governance.\n\n"
                "## Facts, Assumptions, and Validation Questions\n\n"
                "Questions.\n\n"
                "## Workshop Prep\n\n"
                "Prep.\n\n"
                "## Recommended Next Steps\n\n"
                "Next steps.\n\n"
                "## Appendix: Source Ledger\n\n"
                "| Source | Evidence |\n"
                "| --- | --- |\n"
                "| Inventory | Fixture |\n",
                encoding="utf-8",
            )

            render = subprocess.run(
                [
                    sys.executable,
                    str(RENDER_SCRIPT),
                    str(markdown),
                    str(docx),
                    "--title",
                    "Fixture Brief",
                ],
                capture_output=True,
                text=True,
                check=False,
            )
            self.assertEqual(render.returncode, 0, render.stderr)

            verify = subprocess.run(
                [sys.executable, str(VERIFY_SCRIPT), str(docx)],
                capture_output=True,
                text=True,
                check=False,
            )
            self.assertEqual(verify.returncode, 0, verify.stderr)

    @unittest.skipUnless(has_python_docx(), "python-docx is not installed")
    def test_verify_docx_fails_when_required_sections_are_missing(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp:
            docx = Path(tmp) / "bad.docx"
            document = Document()
            document.add_heading("Bad Brief", level=1)
            document.add_paragraph("Missing required structure.")
            document.save(docx)

            verify = subprocess.run(
                [sys.executable, str(VERIFY_SCRIPT), str(docx)],
                capture_output=True,
                text=True,
                check=False,
            )

            self.assertEqual(verify.returncode, 1)
            self.assertIn("Missing required headings", verify.stderr)


if __name__ == "__main__":
    unittest.main()
