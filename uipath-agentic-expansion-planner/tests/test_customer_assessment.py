from __future__ import annotations

import copy
import importlib.util
import json
import re
import shutil
import subprocess
import sys
import tempfile
import unittest
from datetime import date, timedelta
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"
GOLDEN = ROOT / "tests" / "fixtures" / "golden" / "noisy"
FORWARD = ROOT / "tests" / "fixtures" / "forward" / "multi_domain"
PROFILER = SCRIPTS / "inventory_profiler.py"
RENDERER = SCRIPTS / "render_customer_assessment.py"
BUILDER = SCRIPTS / "build_customer_assessment.py"

sys.path.insert(0, str(SCRIPTS))
from assessment_contracts import (  # noqa: E402
    CLAIM_TYPES,
    artifact_sha256,
    derive_readiness,
    has_system_changing_action,
    validate_process_map,
    validate_semantic_review,
)
from portfolio_contracts import score_portfolio, validate_portfolio  # noqa: E402
from validate_customer_assessment import (  # noqa: E402
    markdown_word_count,
    validate_customer_assessment,
)


def load_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8"))


def write_json(path: Path, value):
    path.write_text(
        json.dumps(value, indent=2, sort_keys=True, ensure_ascii=True) + "\n",
        encoding="utf-8",
    )


def profile_noisy(root: Path):
    output = root / "profile"
    result = subprocess.run(
        [sys.executable, str(PROFILER), "--input", str(GOLDEN / "inventory.csv"), "--outdir", str(output)],
        capture_output=True,
        text=True,
        check=False,
    )
    if result.returncode != 0:
        raise AssertionError(result.stderr)
    return load_json(output / "inventory_profile.json"), output / "inventory_profile.json"


def orchestration(
    opportunity_id: str,
    process_id: str,
    inventory_id: str,
    process_name: str,
    owner: str,
):
    return {
        "opportunity_id": opportunity_id,
        "process_id": process_id,
        "pattern": "extend_single",
        "existing_automation_ids": [inventory_id],
        "stages": [
            {
                "sequence": 1,
                "phase": "current_state",
                "name": f"Existing {process_name.lower()} automation",
                "role": "existing_automation",
                "inventory_ids": [inventory_id],
                "action": f"Capture the current {process_name.lower()} work and its exceptions.",
                "human_control": "Operations staff resolve routing exceptions.",
            },
            {
                "sequence": 2,
                "phase": "pilot",
                "name": f"tracks {process_name.lower()} waits and handoffs",
                "role": "maestro",
                "inventory_ids": [],
                "action": "Track each sampled case, wait state, and handoff in a read-only pilot log.",
                "human_control": "Finance owners approve process changes.",
            },
            {
                "sequence": 3,
                "phase": "pilot",
                "name": f"outputs {process_name.lower()} summary and flags evidence gaps",
                "role": "agentic",
                "inventory_ids": [],
                "action": "Output one case summary and flag missing evidence for human review.",
                "human_control": "A reviewer approves every proposed action.",
            },
            {
                "sequence": 4,
                "phase": "pilot",
                "name": f"{owner} reviews weekly",
                "role": "human_review",
                "inventory_ids": [],
                "action": "Compare each output with the historical result in a weekly review.",
                "human_control": f"The {owner} owns the final decision.",
            },
            {
                "sequence": 5,
                "phase": "future_state",
                "name": "Record human-approved transaction",
                "role": "robot",
                "inventory_ids": [],
                "action": "Record the human-approved update in the existing finance system.",
                "human_control": "Only a human-approved transaction may reach the production system.",
            },
        ],
        "capability_roles": {
            "maestro": {
                "applicability": "applies",
                "role": f"Coordinates {process_name.lower()} work, approvals, wait states, and exceptions.",
            },
            "agentic": {
                "applicability": "applies",
                "role": f"Assembles {process_name.lower()} context and recommends the next action for review.",
            },
            "genai": {
                "applicability": "not_needed",
                "role": "The first phase can use structured case summaries without generated customer content.",
            },
            "robots": {
                "applicability": "applies",
                "role": "Execute approved updates through the existing finance automation interfaces.",
            },
            "human_review": {
                "applicability": "applies",
                "role": f"The {owner} retains every final resolution and system-update decision.",
            },
        },
        "measurement_plan": {
            "sample_method": (
                "100 recent completed cases selected across the observed exception types"
            ),
            "ground_truth": (
                f"The historical {owner} disposition for each sampled case"
            ),
            "ground_truth_owner": owner,
            "metrics": [
                {
                    "name": "agreement",
                    "formula": "reviewed cases with matching proposed routes / reviewed cases",
                },
                {
                    "name": "coverage",
                    "formula": "cases with complete evidence / reviewed cases",
                },
            ],
            "cadence": "weekly",
            "owner": owner,
            "mixed_result_action": (
                "Correct the rules, rerun failed cases, and do not proceed until every gate passes."
            ),
        },
        "next_step": {
            "owner": owner,
            "account_team_owner": "UiPath customer success manager",
            "action": (
                f"confirm the {process_name.lower()} boundary and 100-case sample; "
                "book the day-21 review"
            ),
            "deliverable": f"an approved {process_name.lower()} pilot charter",
            "target_days": 14,
        },
    }


def process_map(profile, portfolio, validation_date: date):
    process_specs = [
        (
            "PROC-INVOICE-RESOLUTION",
            "Invoice exception resolution",
            "Finance operations",
            "INV-INVENTORY-R00002",
            "an invoice exception fails the approved validation rules",
            "a finance reviewer approves the resolution and SAP records the result",
            "suppliers receive a traceable and timely payment decision",
            "Finance operations lead",
        ),
        (
            "PROC-VENDOR-SUBMISSION",
            "Vendor submission coordination",
            "Vendor operations",
            "INV-INVENTORY-R00004",
            "a vendor submission arrives in the shared intake channel",
            "vendor operations routes a complete case to the finance owner",
            "complete submissions reach the correct owner without manual inbox triage",
            "Vendor operations lead",
        ),
        (
            "PROC-FINANCE-EVIDENCE",
            "Finance evidence preparation",
            "Finance assurance",
            "INV-INVENTORY-R00005",
            "an assurance evidence request is logged for finance",
            "the assurance owner accepts a complete evidence pack",
            "reviewers receive traceable evidence with a human-owned disposition",
            "Finance assurance lead",
        ),
    ]
    opportunity_ids = portfolio["rankings"]["high_impact"][:3]
    selected_processes = {
        process_specs[index][0]: opportunity_id
        for index, opportunity_id in enumerate(opportunity_ids)
    }
    decisions = []
    for index, (process_id, name, *_rest) in enumerate(process_specs, start=1):
        selected = process_id in selected_processes
        decisions.append(
            {
                "process_id": process_id,
                "status": "selected" if selected else "deferred",
                "rank": index,
                "strategy_alignment": "confirmed" if selected else "validation_required",
                "rationale": (
                    f"Selected because {name.lower()} has stronger source support and a bounded pilot ahead of the remaining mapped processes."
                    if selected
                    else f"Defer because higher-ranked pilots have stronger evidence until the owner confirms this process baseline and customer priority."
                ),
                "observed_evidence": [
                    f"The inventory includes one mapped record for {name.lower()}",
                    "The mapped record has a named lifecycle status and business function",
                ],
                "validation_needed": [
                    "Customer owner confirmation of the process boundary and current baseline"
                ],
            }
        )
    return {
        "schema_version": "1.0",
        "process_map_id": "PROCESS-MAP-NOISY-001",
        "customer_name": portfolio["customer_name"],
        "profile_source_sha256": profile["metadata"]["source_sha256"],
        "confirmation_status": "analyst_confirmed",
        "confirmed_by": "Account planning analyst",
        "confirmed_on": validation_date.isoformat(),
        "prioritization": {
            "method": (
                "Rank confirmed processes by strategy alignment, workload signal, automation "
                "foundation, evidence quality, and delivery risk."
            ),
            "criteria": [
                "strategy_alignment",
                "workload_signal",
                "automation_foundation",
                "evidence_quality",
                "delivery_risk",
            ],
            "decisions": decisions,
        },
        "processes": [
            {
                "process_id": process_id,
                "name": name,
                "business_function": business_function,
                "boundary": {
                    "starts_when": starts_when,
                    "ends_when": ends_when,
                    "business_outcome": business_outcome,
                },
                "inventory_ids": [inventory_id],
                "membership_rationale": f"The mapped automation directly supports the {name.lower()} outcome and boundary.",
                "linkage": {
                    "status": "not_applicable",
                    "rationale": "The process contains one mapped automation record only.",
                    "validation_step": "No cross-record case linkage is required for this process.",
                },
            }
            for (
                process_id,
                name,
                business_function,
                inventory_id,
                starts_when,
                ends_when,
                business_outcome,
                _owner,
            ) in process_specs
        ],
        "unmapped_inventory": [
            {
                "inventory_id": "INV-INVENTORY-R00003",
                "reason": "The archived duplicate cannot support a distinct current process recommendation.",
            }
        ],
        "orchestrations": [
            orchestration(
                opportunity_id,
                process_specs[index][0],
                process_specs[index][3],
                process_specs[index][1],
                process_specs[index][7],
            )
            for index, opportunity_id in enumerate(opportunity_ids)
        ],
    }


def claim_reviews(
    opportunity: dict,
    process_id: str,
    *,
    fail_claim: str | None = None,
):
    evidence = opportunity["evidence_refs"]
    refs = {
        "inventory_support": evidence["inventory_ids"],
        "strategy_support": evidence["public_source_ids"],
        "process_coherence": [process_id],
        "agentic_need": [evidence["inventory_ids"][0], process_id],
        "capability_fit": [process_id],
        "value_logic": [evidence["inventory_ids"][0], *evidence["assumption_ids"]],
        "pilot_realism": [process_id],
        "customer_language": [],
    }
    values = []
    for claim_type in sorted(CLAIM_TYPES):
        judgment = "fail" if claim_type == fail_claim else "pass"
        if claim_type in {"capability_fit", "value_logic"} and fail_claim is None:
            judgment = "needs_validation"
        values.append(
            {
                "claim_type": claim_type,
                "judgment": judgment,
                "evidence_refs": refs[claim_type],
                "rationale": "The reviewed evidence supports this conclusion at the declared readiness level.",
            }
        )
    return values


def create_case(root: Path, *, opportunity_count: int = 1, reviewer_mode: str = "independent_agent"):
    validation_date = date.today()
    profile, profile_path = profile_noisy(root)
    ledger = load_json(GOLDEN / "evidence_ledger.json")
    portfolio = load_json(GOLDEN / "portfolio.json")
    base = portfolio["opportunities"][0]
    opportunities = []
    names = [
        "Invoice exception resolution",
        "Vendor submission coordination",
        "Finance evidence preparation",
    ]
    inventory_ids = [
        "INV-INVENTORY-R00002",
        "INV-INVENTORY-R00004",
        "INV-INVENTORY-R00005",
    ]
    why_now = [
        "Invoice exceptions delay supplier payment decisions and already have a deployed triage foundation.",
        "Vendor submissions still depend on shared-inbox routing before finance can begin review.",
        "Finance assurance requests lack a consistent evidence pack and accountable completion point.",
    ]
    value_levers = [
        "supplier payment decision quality",
        "vendor submission routing time",
        "assurance evidence completeness",
    ]
    owners = [
        "Finance operations lead",
        "Vendor operations lead",
        "Finance assurance lead",
    ]
    for index in range(opportunity_count):
        item = copy.deepcopy(base)
        item["opportunity_id"] = f"OPP-CUSTOMER-{index + 1:02d}"
        item["name"] = names[index]
        item["why_now"] = why_now[index]
        item["value_levers"] = [value_levers[index]]
        item["evidence_refs"]["inventory_ids"] = list(
            dict.fromkeys(["INV-INVENTORY-R00002", inventory_ids[index]])
        )
        item["pilot"]["exit_criteria"][0] = (
            "Stop if agreement is under 80% or any control breach occurs. "
            "Go if 100 cases reach 95% reviewer agreement. "
            "Revise if reviewer agreement is between 80-94%."
        )
        item["pilot"]["data_needed"][0] = "100 completed historical cases with reviewer outcomes"
        item["pilot"]["narrow_scope"] = (
            "Historical read-only shadow of 100 completed cases with no source-system changes."
        )
        item["pilot"]["owner"] = owners[index]
        item["pilot"]["timeline_days"] = 21
        item["criteria_scores"]["strategic_alignment"] = 5 - index
        item["criteria_scores"]["time_to_pilot"] = 4 - index
        opportunities.append(item)
    portfolio["opportunities"] = opportunities
    portfolio["ranking_limits"] = {
        "high_impact": opportunity_count,
        "low_friction_poc": opportunity_count,
    }
    portfolio = score_portfolio(portfolio)
    mapping = process_map(profile, portfolio, validation_date)

    ledger_path = root / "evidence_ledger.json"
    portfolio_path = root / "portfolio.json"
    process_map_path = root / "process_map.json"
    write_json(ledger_path, ledger)
    write_json(portfolio_path, portfolio)
    write_json(process_map_path, mapping)
    hashes = {
        "inventory_profile_sha256": artifact_sha256(profile_path),
        "evidence_ledger_sha256": artifact_sha256(ledger_path),
        "portfolio_sha256": artifact_sha256(portfolio_path),
        "process_map_sha256": artifact_sha256(process_map_path),
    }
    review = {
        "schema_version": "1.0",
        "review_id": "REVIEW-NOISY-001",
        "portfolio_id": portfolio["portfolio_id"],
        "process_map_id": mapping["process_map_id"],
        "reviewed_at": validation_date.isoformat(),
        "reviewer": {"mode": reviewer_mode, "id": "independent-review-fixture"},
        "artifact_hashes": hashes,
        "opportunity_reviews": [
            {
                "opportunity_id": opportunity_id,
                "claim_reviews": claim_reviews(
                    next(
                        item
                        for item in portfolio["opportunities"]
                        if item["opportunity_id"] == opportunity_id
                    ),
                    next(
                        item["process_id"]
                        for item in mapping["orchestrations"]
                        if item["opportunity_id"] == opportunity_id
                    ),
                ),
                "blocking_findings": [],
            }
            for opportunity_id in portfolio["rankings"]["high_impact"][:3]
        ],
        "overall_readiness": "workshop_ready" if reviewer_mode != "single_agent_fallback" else "exploratory",
    }
    review_path = root / "semantic_review.json"
    write_json(review_path, review)
    return {
        "date": validation_date,
        "profile": profile,
        "profile_path": profile_path,
        "ledger": ledger,
        "ledger_path": ledger_path,
        "portfolio": portfolio,
        "portfolio_path": portfolio_path,
        "process_map": mapping,
        "process_map_path": process_map_path,
        "review": review,
        "review_path": review_path,
        "hashes": hashes,
    }


class CustomerAssessmentTests(unittest.TestCase):
    def test_mutation_detection_is_broad_and_explicitly_negated(self):
        for action in (
            "Create a supplier record.",
            "Delete the source record.",
            "Approved the payment decision.",
            "Submit the case for payment.",
            "Publish the approved decision.",
            "Upload the result to the customer system.",
        ):
            with self.subTest(action=action):
                self.assertTrue(has_system_changing_action(action))
        for action in (
            "Observe historical cases in read-only shadow mode.",
            "Compare evidence with the approved checklist.",
            "Do not create, update, or delete source records.",
            "Compile evidence without changing source systems.",
        ):
            with self.subTest(action=action):
                self.assertFalse(has_system_changing_action(action))

    def test_profile_11_preserves_detailed_lifecycle_and_safe_source_metadata(self):
        statuses = [
            "Deployed",
            "Pilot",
            "Paused",
            "Retired",
            "Cancelled",
            "Rejected",
            "Duplicate",
            "Idea",
            "Unknown",
            "Inactive",
        ]
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / "portfolio.csv"
            rows = ["Name,Description,Status,Department"]
            rows.extend(
                f"Case {index},A complete process description for case {index},{status},Operations"
                for index, status in enumerate(statuses, start=1)
            )
            source.write_text("\n".join(rows) + "\n", encoding="utf-8")
            result = subprocess.run(
                [sys.executable, str(PROFILER), "--input", str(source), "--outdir", str(root / "out")],
                capture_output=True,
                text=True,
                check=False,
            )
            self.assertEqual(result.returncode, 0, result.stderr)
            profile = load_json(root / "out" / "inventory_profile.json")
            self.assertEqual(profile["schema_version"], "1.1")
            self.assertEqual(profile["metadata"]["source_name"], "portfolio.csv")
            self.assertEqual(profile["metadata"]["source_sha256"], artifact_sha256(source))
            counts = profile["status_summary"]["lifecycle_status_counts"]
            for status in (
                "deployed",
                "pipeline",
                "paused",
                "retired",
                "cancelled",
                "rejected",
                "duplicate",
                "idea",
                "unknown",
                "other",
            ):
                self.assertEqual(counts[status], 1)

    def test_profile_splits_multi_system_cells_into_concentration_counts(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / "portfolio.csv"
            source.write_text(
                "Name,Description,Status,Department,Systems\n"
                "Case 1,First process,Deployed,Finance,SAP; Outlook\n"
                "Case 2,Second process,Pipeline,Finance,SAP | ServiceNow\n"
                "Case 3,Third process,Idea,IT,Outlook\n",
                encoding="utf-8",
            )
            result = subprocess.run(
                [sys.executable, str(PROFILER), "--input", str(source), "--outdir", str(root / "out")],
                capture_output=True,
                text=True,
                check=False,
            )
            self.assertEqual(result.returncode, 0, result.stderr)
            profile = load_json(root / "out" / "inventory_profile.json")
            self.assertEqual(
                profile["systems_summary"]["systems_top_values"],
                [
                    {"value": "SAP", "count": 2},
                    {"value": "Outlook", "count": 2},
                    {"value": "ServiceNow", "count": 1},
                ],
            )

    def test_customer_concentration_discloses_cutoff_ties(self):
        from render_customer_assessment import concentration_text

        departments = [
            "Information Technology",
            "Information Technology",
            "Information Technology",
            "Procurement",
            "Procurement",
            "Finance",
            "Legal",
            "Programs",
            "Human Resources",
            "Shared Services",
        ]
        profile = {
            "inventory_items": [
                {"department": value, "lifecycle_status": "deployed"}
                for value in departments
            ]
        }
        summary = concentration_text(profile, "department", "departments")
        self.assertEqual(
            summary,
            "Information Technology (3); Procurement (2); "
            "5 departments tied at 1 each",
        )

    def test_profile_10_remains_valid_for_legacy_portfolio_cross_check(self):
        with tempfile.TemporaryDirectory() as tmp:
            case = create_case(Path(tmp))
            legacy = copy.deepcopy(case["profile"])
            legacy["schema_version"] = "1.0"
            legacy["metadata"].pop("source_name")
            legacy["metadata"].pop("source_sha256")
            legacy["status_summary"].pop("lifecycle_status_counts")
            for item in legacy["inventory_items"]:
                item.pop("raw_status")
                item.pop("lifecycle_status")
            self.assertEqual(
                validate_portfolio(case["portfolio"], case["ledger"], profile=legacy), []
            )

    def test_profile_11_rejects_tampered_footprint_counts_and_duplicate_ids(self):
        with tempfile.TemporaryDirectory() as tmp:
            case = create_case(Path(tmp))
            wrong_counts = copy.deepcopy(case["profile"])
            wrong_counts["status_summary"]["lifecycle_status_counts"]["deployed"] += 1
            self.assertTrue(
                any(
                    "lifecycle_status_counts must exactly match" in item
                    for item in validate_portfolio(
                        case["portfolio"], case["ledger"], profile=wrong_counts
                    )
                )
            )
            duplicate_id = copy.deepcopy(case["profile"])
            duplicate_id["inventory_items"][1]["inventory_id"] = duplicate_id[
                "inventory_items"
            ][0]["inventory_id"]
            self.assertTrue(
                any(
                    "duplicate ID" in item
                    for item in validate_portfolio(
                        case["portfolio"], case["ledger"], profile=duplicate_id
                    )
                )
            )
            missing_id = copy.deepcopy(case["profile"])
            missing_id["inventory_items"][0].pop("inventory_id")
            self.assertTrue(
                any(
                    "inventory_id must match" in item
                    for item in validate_portfolio(
                        case["portfolio"], case["ledger"], profile=missing_id
                    )
                )
            )
            invalid_sheet_columns = copy.deepcopy(case["profile"])
            invalid_sheet_columns["sheets"][0]["column_count"] += 1
            self.assertTrue(
                any(
                    "column_count must equal columns length" in item
                    for item in validate_portfolio(
                        case["portfolio"], case["ledger"], profile=invalid_sheet_columns
                    )
                )
            )
            missing_sheet_mapping = copy.deepcopy(case["profile"])
            missing_sheet_mapping.pop("sheet_field_mappings")
            self.assertTrue(
                any(
                    "sheet_field_mappings must be an object" in item
                    for item in validate_portfolio(
                        case["portfolio"], case["ledger"], profile=missing_sheet_mapping
                    )
                )
            )
            leaked_source_path = copy.deepcopy(case["profile"])
            leaked_source_path["metadata"]["source_file"] = (
                "/" + "Users/example/customer/inventory.csv"
            )
            self.assertTrue(
                any(
                    "source_file must be a display-safe basename" in item
                    for item in validate_portfolio(
                        case["portfolio"], case["ledger"], profile=leaked_source_path
                    )
                )
            )
            invalid_coverage = copy.deepcopy(case["profile"])
            invalid_coverage["data_quality"]["field_coverage"]["status"][
                "coverage_pct"
            ] = 1
            self.assertTrue(
                any(
                    "coverage_pct must match nonblank and total rows" in item
                    for item in validate_portfolio(
                        case["portfolio"], case["ledger"], profile=invalid_coverage
                    )
                )
            )

    def test_profile_source_date_rejects_stale_ledger_claim_and_tampering(self):
        profile = load_json(FORWARD / "profile" / "inventory_profile.json")
        ledger = load_json(FORWARD / "evidence_ledger.json")
        portfolio = load_json(FORWARD / "portfolio.json")
        self.assertEqual(validate_portfolio(portfolio, ledger, profile=profile), [])

        stale_claim = copy.deepcopy(ledger)
        stale_claim["inventory_profile"]["as_of_date"] = "2026-07-21"
        self.assertTrue(
            any(
                "must equal the latest valid source record date" in item
                for item in validate_portfolio(portfolio, stale_claim, profile=profile)
            )
        )

        tampered_profile = copy.deepcopy(profile)
        tampered_profile["metadata"]["source_date_summary"]["latest_date"] = (
            "2026-07-21"
        )
        self.assertTrue(
            any(
                "latest_date must match inventory_items" in item
                for item in validate_portfolio(
                    portfolio, ledger, profile=tampered_profile
                )
            )
        )

    def test_process_map_requires_complete_non_overlapping_confirmed_mapping(self):
        with tempfile.TemporaryDirectory() as tmp:
            case = create_case(Path(tmp))
            self.assertEqual(
                validate_process_map(case["process_map"], case["profile"], case["portfolio"]),
                [],
            )
            missing = copy.deepcopy(case["process_map"])
            missing["unmapped_inventory"] = []
            self.assertTrue(
                any("does not cover" in item for item in validate_process_map(missing, case["profile"], case["portfolio"]))
            )
            duplicate = copy.deepcopy(case["process_map"])
            duplicate["processes"][0]["inventory_ids"].append("INV-INVENTORY-R00005")
            self.assertTrue(
                any("assigned more than once" in item for item in validate_process_map(duplicate, case["profile"], case["portfolio"]))
            )
            bad_boundary = copy.deepcopy(case["process_map"])
            bad_boundary["processes"][0]["boundary"]["starts_when"] = "invoice"
            self.assertTrue(
                any("at least 4 words" in item for item in validate_process_map(bad_boundary, case["profile"], case["portfolio"]))
            )
            generic_boundary = copy.deepcopy(case["process_map"])
            generic_boundary["processes"][0]["boundary"][
                "ends_when"
            ] = "the approved result reaches the responsible system"
            self.assertTrue(
                any(
                    "generic placeholder language" in item
                    for item in validate_process_map(
                        generic_boundary, case["profile"], case["portfolio"]
                    )
                )
            )
            placeholder_confirmation = copy.deepcopy(case["process_map"])
            placeholder_confirmation["confirmed_by"] = "TBD"
            self.assertTrue(
                any(
                    "confirming analyst" in item
                    for item in validate_process_map(
                        placeholder_confirmation, case["profile"], case["portfolio"]
                    )
                )
            )
            placeholder_account_owner = copy.deepcopy(case["process_map"])
            placeholder_account_owner["orchestrations"][0]["next_step"][
                "account_team_owner"
            ] = "TBD"
            self.assertTrue(
                any(
                    "account_team_owner must be an accountable" in item
                    for item in validate_process_map(
                        placeholder_account_owner, case["profile"], case["portfolio"]
                    )
                )
            )
            missing_stage = copy.deepcopy(case["process_map"])
            missing_stage["orchestrations"][0]["stages"] = [
                stage
                for stage in missing_stage["orchestrations"][0]["stages"]
                if stage["role"] != "maestro"
            ]
            for index, stage in enumerate(missing_stage["orchestrations"][0]["stages"], start=1):
                stage["sequence"] = index
            self.assertTrue(
                any(
                    "no maestro stage" in item
                    for item in validate_process_map(missing_stage, case["profile"], case["portfolio"])
                )
            )
            contradictory_role = copy.deepcopy(case["process_map"])
            contradictory_role["orchestrations"][0]["capability_roles"]["maestro"][
                "applicability"
            ] = "not_needed"
            self.assertTrue(
                any(
                    "says not_needed" in item
                    for item in validate_process_map(
                        contradictory_role, case["profile"], case["portfolio"]
                    )
                )
            )
            missing_linkage = copy.deepcopy(case["process_map"])
            missing_linkage["processes"][0].pop("linkage")
            self.assertTrue(
                any(
                    "missing required field(s): linkage" in item
                    for item in validate_process_map(
                        missing_linkage, case["profile"], case["portfolio"]
                    )
                )
            )
            invalid_multi_linkage = copy.deepcopy(case["process_map"])
            invalid_multi_linkage["processes"][0]["inventory_ids"].append(
                "INV-INVENTORY-R00003"
            )
            invalid_multi_linkage["unmapped_inventory"] = []
            self.assertTrue(
                any(
                    "cannot be not_applicable for a multi-automation process" in item
                    for item in validate_process_map(
                        invalid_multi_linkage, case["profile"], case["portfolio"]
                    )
                )
            )
            unsafe_write_order = copy.deepcopy(case["process_map"])
            stages = unsafe_write_order["orchestrations"][0]["stages"]
            stages[3], stages[4] = stages[4], stages[3]
            for index, stage in enumerate(stages, start=1):
                stage["sequence"] = index
            self.assertTrue(
                any(
                    "before a human-review stage" in item
                    for item in validate_process_map(
                        unsafe_write_order, case["profile"], case["portfolio"]
                    )
                )
            )
            pilot_write = copy.deepcopy(case["process_map"])
            pilot_write["orchestrations"][0]["stages"][4]["phase"] = "pilot"
            self.assertTrue(
                any(
                    "system-changing action must be future_state" in item
                    for item in validate_process_map(
                        pilot_write, case["profile"], case["portfolio"]
                    )
                )
            )
            read_only_bypass = copy.deepcopy(case["process_map"])
            read_only_bypass["orchestrations"][0]["stages"][2]["action"] = (
                "Delete source records while operating in read-only shadow mode."
            )
            self.assertTrue(
                any(
                    "system-changing action must be future_state" in item
                    for item in validate_process_map(
                        read_only_bypass, case["profile"], case["portfolio"]
                    )
                )
            )
            explicitly_negated_write = copy.deepcopy(case["process_map"])
            explicitly_negated_write["orchestrations"][0]["stages"][2]["action"] = (
                "Output a historical case summary in weekly shadow mode; do not delete or update source records."
            )
            self.assertEqual(
                validate_process_map(
                    explicitly_negated_write, case["profile"], case["portfolio"]
                ),
                [],
            )
            ambiguous_write = copy.deepcopy(case["process_map"])
            write_stage = ambiguous_write["orchestrations"][0]["stages"][4]
            write_stage["name"] = "Record approved status"
            write_stage["action"] = "Record the approved status in the finance system."
            write_stage["human_control"] = "An approval gates the production update."
            self.assertTrue(
                any(
                    "human-confirmed or human-approved" in item
                    for item in validate_process_map(
                        ambiguous_write, case["profile"], case["portfolio"]
                    )
                )
            )
            generic_gate = copy.deepcopy(case["portfolio"])
            generic_gate["opportunities"][0]["pilot"]["exit_criteria"][0] = (
                "Stop if a control fails. Go if the agreed sample and quality conditions pass. "
                "Revise if every other result occurs."
            )
            self.assertTrue(
                any(
                    "generic phrase" in item or "quantitative" in item
                    for item in validate_process_map(
                        case["process_map"], case["profile"], generic_gate
                    )
                )
            )
            uncovered_failure_range = copy.deepcopy(case["portfolio"])
            uncovered_failure_range["opportunities"][0]["pilot"]["exit_criteria"][0] = (
                "Stop if a control breach or prohibited automated action occurs. "
                "Go if 100 cases achieve at least 95 percent reviewer agreement. "
                "Revise if reviewer agreement is 80 to 94 percent."
            )
            self.assertTrue(
                any(
                    "quantitative failure range" in item
                    for item in validate_process_map(
                        case["process_map"], case["profile"], uncovered_failure_range
                    )
                )
            )
            impossible_shadow_metric = copy.deepcopy(case["portfolio"])
            impossible_shadow_metric["opportunities"][0]["pilot"]["narrow_scope"] = (
                "Replay historical cases in read-only shadow mode."
            )
            impossible_shadow_metric["opportunities"][0]["pilot"]["exit_criteria"][0] = (
                "Stop if a control breach occurs or measured cycle-time reduction is below 10 percent. "
                "Go if 100 cases show at least 20 percent lower cycle time. "
                "Revise if cycle-time reduction is 10 to 19 percent."
            )
            self.assertTrue(
                any(
                    "cannot measure live cycle-time reduction" in item
                    for item in validate_process_map(
                        case["process_map"], case["profile"], impossible_shadow_metric
                    )
                )
            )
            mismatched_sample = copy.deepcopy(case["portfolio"])
            mismatched_sample["opportunities"][0]["pilot"]["data_needed"][0] = (
                "20 completed historical cases with reviewer outcomes"
            )
            self.assertTrue(
                any(
                    "numeric sample used by the proceed gate" in item
                    for item in validate_process_map(
                        case["process_map"], case["profile"], mismatched_sample
                    )
                )
            )
            missing_cadence = copy.deepcopy(case["process_map"])
            for stage in missing_cadence["orchestrations"][0]["stages"]:
                if stage["phase"] == "pilot":
                    stage["action"] = re.sub(
                        r"\b(?:daily|each case|every case|per case|weekly)\b",
                        "during the pilot",
                        stage["action"],
                        flags=re.I,
                    )
            self.assertTrue(
                any(
                    "observation cadence" in item
                    for item in validate_process_map(
                        missing_cadence, case["profile"], case["portfolio"]
                    )
                )
            )
            owner_mismatch = copy.deepcopy(case["process_map"])
            owner_mismatch["orchestrations"][0]["next_step"]["owner"] = (
                "Unrelated program owner"
            )
            self.assertTrue(
                any(
                    "must match the pilot decision owner" in item
                    for item in validate_process_map(
                        owner_mismatch, case["profile"], case["portfolio"]
                    )
                )
            )
            missing_ground_truth = copy.deepcopy(case["process_map"])
            missing_ground_truth["orchestrations"][0]["measurement_plan"].pop(
                "ground_truth"
            )
            self.assertTrue(
                any(
                    "missing required field(s): ground_truth" in item
                    for item in validate_process_map(
                        missing_ground_truth, case["profile"], case["portfolio"]
                    )
                )
            )
            missing_ground_truth_owner = copy.deepcopy(case["process_map"])
            missing_ground_truth_owner["orchestrations"][0]["measurement_plan"].pop(
                "ground_truth_owner"
            )
            self.assertTrue(
                any(
                    "missing required field(s): ground_truth_owner" in item
                    for item in validate_process_map(
                        missing_ground_truth_owner,
                        case["profile"],
                        case["portfolio"],
                    )
                )
            )
            bad_formula = copy.deepcopy(case["process_map"])
            bad_formula["orchestrations"][0]["measurement_plan"]["metrics"][0][
                "formula"
            ] = "reviewer agreement percentage"
            self.assertTrue(
                any(
                    "must state a numerator and denominator" in item
                    for item in validate_process_map(
                        bad_formula, case["profile"], case["portfolio"]
                    )
                )
            )
            mismatched_ratio_units = copy.deepcopy(case["process_map"])
            mismatched_ratio_units["orchestrations"][0]["measurement_plan"]["metrics"][0][
                "formula"
            ] = "complete required evidence fields / reviewed cases"
            self.assertTrue(
                any(
                    "comparable numerator and denominator units" in item
                    for item in validate_process_map(
                        mismatched_ratio_units,
                        case["profile"],
                        case["portfolio"],
                    )
                )
            )
            measurement_owner_mismatch = copy.deepcopy(case["process_map"])
            measurement_owner_mismatch["orchestrations"][0]["measurement_plan"][
                "owner"
            ] = "Unrelated analyst"
            self.assertTrue(
                any(
                    "measurement_plan.owner must match" in item
                    for item in validate_process_map(
                        measurement_owner_mismatch,
                        case["profile"],
                        case["portfolio"],
                    )
                )
            )
            missing_rerun = copy.deepcopy(case["process_map"])
            missing_rerun["orchestrations"][0]["measurement_plan"][
                "mixed_result_action"
            ] = "Correct the rules and wait for the owner before proceeding further."
            self.assertTrue(
                any(
                    "must require a rerun" in item
                    for item in validate_process_map(
                        missing_rerun, case["profile"], case["portfolio"]
                    )
                )
            )
            unexplained_defer = copy.deepcopy(case["process_map"])
            decision = unexplained_defer["prioritization"]["decisions"][-1]
            decision["strategy_alignment"] = "confirmed"
            decision["rationale"] = "This process remains a useful follow-on opportunity for the customer portfolio."
            self.assertTrue(
                any(
                    "why and until when" in item
                    for item in validate_process_map(
                        unexplained_defer, case["profile"], case["portfolio"]
                    )
                )
            )
            missing_observed_evidence = copy.deepcopy(case["process_map"])
            missing_observed_evidence["prioritization"]["decisions"][0][
                "observed_evidence"
            ] = []
            self.assertTrue(
                any(
                    "observed_evidence" in item
                    for item in validate_process_map(
                        missing_observed_evidence, case["profile"], case["portfolio"]
                    )
                )
            )
            mixed_evidence = copy.deepcopy(case["process_map"])
            mixed_evidence["prioritization"]["decisions"][0][
                "observed_evidence"
            ][0] = "Validate the assumed current approval-history coverage"
            self.assertTrue(
                any(
                    "mixes observed facts" in item
                    for item in validate_process_map(
                        mixed_evidence, case["profile"], case["portfolio"]
                    )
                )
            )
            weak_selection = copy.deepcopy(case["process_map"])
            weak_selection["prioritization"]["decisions"][0]["rationale"] = (
                "This process appears useful for the customer portfolio and deserves attention."
            )
            self.assertTrue(
                any(
                    "selected over alternatives" in item
                    for item in validate_process_map(
                        weak_selection, case["profile"], case["portfolio"]
                    )
                )
            )
            outside_evidence = copy.deepcopy(case["process_map"])
            outside_evidence["orchestrations"][0]["existing_automation_ids"] = [
                "INV-INVENTORY-R00003"
            ]
            outside_evidence["orchestrations"][0]["pattern"] = "extend_single"
            self.assertTrue(
                any(
                    "opportunity evidence" in item
                    for item in validate_process_map(outside_evidence, case["profile"], case["portfolio"])
                )
            )
            missing_decision = copy.deepcopy(case["process_map"])
            missing_decision["prioritization"]["decisions"] = missing_decision[
                "prioritization"
            ]["decisions"][:-1]
            self.assertTrue(
                any(
                    "prioritization does not cover" in item
                    for item in validate_process_map(
                        missing_decision, case["profile"], case["portfolio"]
                    )
                )
            )
            reversed_selection = copy.deepcopy(case["process_map"])
            reversed_selection["prioritization"]["decisions"][0]["rank"] = 2
            reversed_selection["prioritization"]["decisions"][1]["rank"] = 1
            self.assertTrue(
                any(
                    "selected process prioritization ranks" in item
                    for item in validate_process_map(
                        reversed_selection, case["profile"], case["portfolio"]
                    )
                )
            )

    def test_semantic_review_schema_does_not_force_agentic_use(self):
        schema = load_json(ROOT / "references" / "contracts" / "semantic_review.v1.schema.json")
        claim_type = schema["$defs"]["claimReview"]["properties"]["claim_type"]
        guidance = claim_type["description"].casefold()
        self.assertIn("supports either applies or not_needed", guidance)
        self.assertIn("does not require an agent", guidance)

    def test_semantic_review_rejects_tampering_staleness_and_overstated_fallback(self):
        with tempfile.TemporaryDirectory() as tmp:
            case = create_case(Path(tmp))
            self.assertEqual(
                validate_semantic_review(
                    case["review"],
                    case["ledger"],
                    case["portfolio"],
                    case["process_map"],
                    case["profile"],
                    expected_hashes=case["hashes"],
                    today=case["date"],
                    required_readiness="workshop_ready",
                ),
                [],
            )
            tampered = copy.deepcopy(case["review"])
            tampered["artifact_hashes"]["portfolio_sha256"] = "0" * 64
            self.assertTrue(
                any(
                    "hash mismatch" in item
                    for item in validate_semantic_review(
                        tampered,
                        case["ledger"],
                        case["portfolio"],
                        case["process_map"],
                        case["profile"],
                        expected_hashes=case["hashes"],
                        today=case["date"],
                    )
                )
            )
            missing_claim = copy.deepcopy(case["review"])
            missing_claim["opportunity_reviews"][0]["claim_reviews"] = missing_claim[
                "opportunity_reviews"
            ][0]["claim_reviews"][:-1]
            self.assertTrue(
                any(
                    "missing claim review" in item
                    for item in validate_semantic_review(
                        missing_claim,
                        case["ledger"],
                        case["portfolio"],
                        case["process_map"],
                        case["profile"],
                        expected_hashes=case["hashes"],
                        today=case["date"],
                    )
                )
            )
            unrelated_evidence = copy.deepcopy(case["review"])
            inventory_review = next(
                item
                for item in unrelated_evidence["opportunity_reviews"][0]["claim_reviews"]
                if item["claim_type"] == "inventory_support"
            )
            inventory_review["evidence_refs"] = ["INV-INVENTORY-R00004"]
            self.assertTrue(
                any(
                    "inventory outside the recommendation" in item
                    for item in validate_semantic_review(
                        unrelated_evidence,
                        case["ledger"],
                        case["portfolio"],
                        case["process_map"],
                        case["profile"],
                        expected_hashes=case["hashes"],
                        today=case["date"],
                    )
                )
            )
            stale = copy.deepcopy(case["review"])
            stale["reviewed_at"] = (case["date"] - timedelta(days=31)).isoformat()
            self.assertTrue(
                any(
                    "days old" in item or "predates" in item
                    for item in validate_semantic_review(
                        stale,
                        case["ledger"],
                        case["portfolio"],
                        case["process_map"],
                        case["profile"],
                        expected_hashes=case["hashes"],
                        today=case["date"],
                    )
                )
            )
            future_evidence = copy.deepcopy(case["ledger"])
            future_evidence["public_sources"][0]["accessed_date"] = (
                case["date"] + timedelta(days=1)
            ).isoformat()
            self.assertTrue(
                any(
                    "predates a bound artifact" in item
                    for item in validate_semantic_review(
                        case["review"],
                        future_evidence,
                        case["portfolio"],
                        case["process_map"],
                        case["profile"],
                        expected_hashes=case["hashes"],
                        today=case["date"],
                    )
                )
            )
            fallback = copy.deepcopy(case["review"])
            fallback["reviewer"]["mode"] = "single_agent_fallback"
            fallback["overall_readiness"] = "workshop_ready"
            self.assertEqual(
                derive_readiness(fallback, case["ledger"], case["portfolio"], case["process_map"]),
                "exploratory",
            )
            self.assertTrue(
                any(
                    "derived readiness" in item
                    for item in validate_semantic_review(
                        fallback,
                        case["ledger"],
                        case["portfolio"],
                        case["process_map"],
                        case["profile"],
                        expected_hashes=case["hashes"],
                        today=case["date"],
                    )
                )
            )

    def test_failed_critical_claim_blocks_workshop_readiness(self):
        with tempfile.TemporaryDirectory() as tmp:
            case = create_case(Path(tmp))
            changed = copy.deepcopy(case["review"])
            opportunity_id = changed["opportunity_reviews"][0]["opportunity_id"]
            opportunity = next(
                item
                for item in case["portfolio"]["opportunities"]
                if item["opportunity_id"] == opportunity_id
            )
            process_id = next(
                item["process_id"]
                for item in case["process_map"]["orchestrations"]
                if item["opportunity_id"] == opportunity_id
            )
            changed["opportunity_reviews"][0]["claim_reviews"] = claim_reviews(
                opportunity,
                process_id,
                fail_claim="process_coherence",
            )
            self.assertEqual(
                derive_readiness(changed, case["ledger"], case["portfolio"], case["process_map"]),
                "exploratory",
            )
            suggested = copy.deepcopy(case["process_map"])
            suggested["confirmation_status"] = "suggested"
            self.assertEqual(
                derive_readiness(case["review"], case["ledger"], case["portfolio"], suggested),
                "exploratory",
            )
            no_agent_needed = copy.deepcopy(case["process_map"])
            orchestration = no_agent_needed["orchestrations"][0]
            orchestration["stages"] = [
                stage for stage in orchestration["stages"] if stage["role"] != "agentic"
            ]
            for index, stage in enumerate(orchestration["stages"], start=1):
                stage["sequence"] = index
            orchestration["capability_roles"]["agentic"] = {
                "applicability": "not_needed",
                "role": "Approved rules and human review cover the initial process boundary.",
            }
            self.assertEqual(
                validate_process_map(no_agent_needed, case["profile"], case["portfolio"]),
                [],
            )
            self.assertEqual(
                derive_readiness(
                    case["review"], case["ledger"], case["portfolio"], no_agent_needed
                ),
                "workshop_ready",
            )

    def test_pilot_authorizable_requires_all_claims_assumptions_and_capabilities(self):
        with tempfile.TemporaryDirectory() as tmp:
            case = create_case(Path(tmp))
            ledger = copy.deepcopy(case["ledger"])
            portfolio = copy.deepcopy(case["portfolio"])
            review = copy.deepcopy(case["review"])
            for assumption in ledger["assumptions"]:
                assumption["status"] = "validated"
            entitlement_ref = "ASM-INTEGRATION-ENTITLEMENT"
            for entitlement in ledger["customer"]["entitlements"]:
                entitlement["status"] = "confirmed"
                entitlement["evidence_refs"] = [entitlement_ref]
            for opportunity in portfolio["opportunities"]:
                for fit in opportunity["capability_fit"]:
                    fit["claim"] = "confirmed_entitlement"
                    fit["entitlement_evidence_refs"] = [entitlement_ref]
            for opportunity_review in review["opportunity_reviews"]:
                for claim in opportunity_review["claim_reviews"]:
                    claim["judgment"] = "pass"
            self.assertEqual(validate_portfolio(portfolio, ledger, profile=case["profile"]), [])
            self.assertEqual(
                derive_readiness(review, ledger, portfolio, case["process_map"]),
                "pilot_authorizable",
            )
            unresolved_capability = copy.deepcopy(case["process_map"])
            unresolved_capability["orchestrations"][0]["capability_roles"]["maestro"][
                "applicability"
            ] = "validation_required"
            self.assertEqual(
                derive_readiness(review, ledger, portfolio, unresolved_capability),
                "workshop_ready",
            )
            missing_measures = copy.deepcopy(portfolio)
            missing_measures["opportunities"][0]["pilot"]["success_metrics"] = []
            self.assertEqual(
                derive_readiness(review, ledger, missing_measures, case["process_map"]),
                "workshop_ready",
            )

    def test_customer_renderer_is_deterministic_concise_and_never_fills(self):
        for opportunity_count in (1, 2, 3):
            with self.subTest(opportunity_count=opportunity_count), tempfile.TemporaryDirectory() as tmp:
                root = Path(tmp)
                case = create_case(root, opportunity_count=opportunity_count)
                outputs = [root / "first.md", root / "second.md"]
                for output in outputs:
                    result = subprocess.run(
                        [
                            sys.executable,
                            str(RENDERER),
                            "--inventory-profile",
                            str(case["profile_path"]),
                            "--evidence-ledger",
                            str(case["ledger_path"]),
                            "--portfolio",
                            str(case["portfolio_path"]),
                            "--process-map",
                            str(case["process_map_path"]),
                            "--semantic-review",
                            str(case["review_path"]),
                            "--output",
                            str(output),
                            "--validation-date",
                            case["date"].isoformat(),
                        ],
                        capture_output=True,
                        text=True,
                        check=False,
                    )
                    self.assertEqual(result.returncode, 0, result.stderr)
                self.assertEqual(outputs[0].read_bytes(), outputs[1].read_bytes())
                text = outputs[0].read_text(encoding="utf-8")
                self.assertEqual(len(re.findall(r"^###\s+", text, flags=re.M)), opportunity_count)
                self.assertLessEqual(markdown_word_count(text), 900)
                self.assertEqual(validate_customer_assessment(text), [])
                self.assertNotRegex(text, r"\b(?:INV|SRC|ASM|OPP|PROC)-")
                self.assertIn("Deployed: 1", text)
                self.assertIn("Pipeline: 1", text)
                self.assertIn("Cancelled: 0", text)
                self.assertIn("Rejected: 0", text)
                self.assertIn("Retired: 1", text)
                self.assertIn(
                    "3 analyst-mapped groups; customer confirmation required", text
                )
                self.assertIn("Invoice exception resolution", text)
                self.assertIn("| Unmapped | 1 record", text)
                self.assertIn("**Information available:** names, descriptions, lifecycle", text)
                self.assertNotIn("volume, annual volume", text.casefold())
                self.assertIn("non-official planning", text)
                self.assertIn("| Process/domain groups |", text)
                self.assertIn("| Department concentration |", text)
                self.assertIn("| System concentration |", text)
                self.assertIn("| Assessment boundary |", text)
                if opportunity_count < 3:
                    self.assertIn("no lower-confidence filler was added", text)
                self.assertIn("Function:", text)
                self.assertIn("Start:", text)
                self.assertIn("End:", text)
                self.assertIn("Outcome:", text)
                self.assertIn("**Decision gate:**", text)
                self.assertIn("Decision owner:", text)
                self.assertIn("| Rank | Process | Why this order |", text)
                self.assertIn("proceed when", text.casefold())
                self.assertIn("adjust when", text.casefold())
                self.assertIn("stop when", text.casefold())
                self.assertIn("pilot continuation only", text.casefold())
                self.assertIn("Input:", text)
                self.assertIn("Ground truth:", text)
                self.assertIn("Ground-truth owner:", text)
                self.assertIn("cases with matching routes / cases", text)
                self.assertIn("Maestro", text)
                self.assertIn("UiPath: CSM", text)
                self.assertIn("(Production)", text)
                self.assertIn("1 automation:", text)
                self.assertIn("product", text.casefold())
                self.assertIn("baseline", text.casefold())
                self.assertIn("**Pilot path:** Proposed", text)
                self.assertIn("Pilot: no writes", text)
                self.assertIn("No valid record-update date", text)
                self.assertNotIn("Inventory date:", text)
                self.assertNotIn(">=", text)
                self.assertNotIn("<=", text)
                self.assertNotIn(".;", text)
                self.assertNotIn("(s)", text)
                if opportunity_count < 3:
                    self.assertIn("ready for workshop validation", text)
                else:
                    self.assertNotIn("ready for workshop validation", text)

    def test_plain_language_gate_rejects_repeated_recommendation_content(self):
        with tempfile.TemporaryDirectory() as tmp:
            case = create_case(Path(tmp), opportunity_count=2)
            from render_customer_assessment import render

            text = render(
                case["profile"],
                case["ledger"],
                case["portfolio"],
                case["process_map"],
                case["review"],
            )
            why_values = re.findall(
                r"(?im)^- \*\*Why it matters:\*\*\s+(.+)$", text
            )
            self.assertEqual(len(why_values), 2)
            repeated = text.replace(why_values[1], why_values[0], 1)
            self.assertTrue(
                any(
                    "repeat the same why it matters text" in item
                    for item in validate_customer_assessment(repeated)
                )
            )

    def test_plain_language_gate_rejects_internal_ids_hype_and_long_sentences(self):
        with tempfile.TemporaryDirectory() as tmp:
            case = create_case(Path(tmp))
            from render_customer_assessment import (
                clause,
                embedded_clause,
                render,
                validation_text,
                why_summary,
            )

            valid = render(case["profile"], case["ledger"], case["portfolio"], case["process_map"])
            self.assertEqual(validate_customer_assessment(valid), [])
            invalid = valid.replace(
                "# Automation portfolio assessment: Lakeview Finance Authority",
                "# game-changing next step using INV-INVENTORY-R00002 and schema_version",
            )
            self.assertNotEqual(invalid, valid)
            failures = validate_customer_assessment(invalid)
            self.assertTrue(any("hype" in item for item in failures))
            self.assertTrue(any("internal evidence" in item for item in failures))
            self.assertTrue(any("internal terminology" in item for item in failures))
            ambiguous_control = valid.replace(
                "Pilot: no writes",
                "Update approved status",
                1,
            )
            self.assertTrue(
                any(
                    "ambiguous decision-write phrase" in item
                    for item in validate_customer_assessment(ambiguous_control)
                )
            )
            unsupported_throughput = valid.replace(
                "**Why it matters:**",
                "**Why it matters:** The pilot covers 4,100 annual applications. ",
                1,
            )
            self.assertTrue(
                any(
                    "pilot throughput" in item
                    for item in validate_customer_assessment(unsupported_throughput)
                )
            )
            structural_overclaim = valid.replace(
                "**Limitations:**",
                "**Limitations:** The file is structurally complete. ",
                1,
            )
            self.assertTrue(
                any(
                    "structural completeness" in item
                    for item in validate_customer_assessment(structural_overclaim)
                )
            )
            cryptic = valid.replace(
                "Maestro, Agents, and Robots proposed",
                "Maestro/Robots: fit to validate",
                1,
            )
            self.assertNotEqual(cryptic, valid)
            self.assertTrue(
                any(
                    "cryptic internal phrasing" in item
                    for item in validate_customer_assessment(cryptic)
                )
            )
            false_confirmation = valid.replace(
                "analyst-mapped groups; customer confirmation required",
                "analyst-confirmed groups",
                1,
            )
            false_confirmation_failures = validate_customer_assessment(
                false_confirmation
            )
            self.assertTrue(
                any(
                    "customer confirmation" in item
                    for item in false_confirmation_failures
                )
            )
            local_path = valid.replace(
                "inventory.csv", "/" + "Users/example/customer/inventory.csv", 1
            )
            self.assertTrue(
                any(
                    "local filesystem path" in item
                    for item in validate_customer_assessment(local_path)
                )
            )
            raw_hash = valid.replace(
                "inventory.csv", "a" * 64 + ".csv", 1
            )
            self.assertTrue(
                any(
                    "raw SHA-256" in item
                    for item in validate_customer_assessment(raw_hash)
                )
            )
            symbolic_threshold = valid.replace(
                "Proceed when",
                "Proceed when >=95% agreement and",
                1,
            )
            self.assertNotEqual(symbolic_threshold, valid)
            self.assertTrue(
                any(
                    "plain threshold language" in item
                    for item in validate_customer_assessment(symbolic_threshold)
                )
            )
            self.assertEqual(clause("Approve the process."), "approve the process")
            self.assertEqual(clause("SAP posts the update."), "SAP posts the update")
            self.assertEqual(
                clause("Supplier Management records the decision."),
                "Supplier Management records the decision",
            )
            self.assertEqual(
                embedded_clause("One Procurement-selected exception pathway."),
                "one Procurement-selected exception pathway",
            )
            self.assertEqual(
                why_summary(
                    "Non-official strategy prioritizes completeness; the inventory lists "
                    "4,100 annual applications."
                ),
                "Non-official strategy prioritizes completeness; the inventory lists 4,100 annual applications.",
            )
            missing_foundation = valid.replace(
                "**Existing automation foundation:**",
                "**Current tools:**",
                1,
            )
            self.assertTrue(
                any(
                    "missing field: Existing automation foundation" in item
                    for item in validate_customer_assessment(missing_foundation)
                )
            )
            missing_footprint = valid.replace(
                "| Process/domain groups |", "| Grouping |", 1
            )
            self.assertTrue(
                any(
                    "missing row: Process/domain groups" in item
                    for item in validate_customer_assessment(missing_footprint)
                )
            )
            incomplete_gate = re.sub(
                r"(?m)^- \*\*Decision gate:\*\*.+$",
                "- **Decision gate:** Measure the result.",
                valid,
                count=1,
            )
            self.assertTrue(
                any(
                    "must define stop, proceed, and adjust" in item
                    for item in validate_customer_assessment(incomplete_gate)
                )
            )
            missing_comparison = valid.replace(
                "| Rank | Process | Why this order |", "", 1
            )
            self.assertTrue(
                any(
                    "rank comparison table" in item
                    for item in validate_customer_assessment(missing_comparison)
                )
            )
            missing_threshold_basis = valid.replace(
                "thresholds from baselines and tolerances",
                "thresholds during the workshop",
                1,
            )
            self.assertTrue(
                any(
                    "threshold-confirmation basis" in item
                    for item in validate_customer_assessment(missing_threshold_basis)
                )
            )
            missing_account_sequence = valid.replace(
                "Account team:", "Account support:", 1
            )
            self.assertTrue(
                any(
                    "account-team execution sequence" in item
                    for item in validate_customer_assessment(missing_account_sequence)
                )
            )
            missing_pilot_mechanics = valid.replace(
                "Pilot mechanics:", "Pilot setup:", 1
            )
            self.assertTrue(
                any(
                    "shared pilot mechanics" in item
                    for item in validate_customer_assessment(missing_pilot_mechanics)
                )
            )
            unsafe_decision_scope = valid.replace(
                "No deployment or investment approval",
                "Approve deployment and investment",
                1,
            )
            self.assertTrue(
                any(
                    "bounded decision scope" in item
                    for item in validate_customer_assessment(unsafe_decision_scope)
                )
            )
            missing_decision_owner = valid.replace(
                "Decision owner:", "Decision support:", 1
            )
            self.assertTrue(
                any(
                    "must name one decision owner" in item
                    for item in validate_customer_assessment(missing_decision_owner)
                )
            )
            deployment_unknown = copy.deepcopy(case["portfolio"]["opportunities"][0])
            deployment_unknown["deployment"]["status"] = "requires_validation"
            self.assertIn(
                "Product, deployment, value unconfirmed",
                validation_text(deployment_unknown, None),
            )

    def test_answer_blind_forward_fixture_is_reproducible_and_actionable(self):
        profile_path = FORWARD / "profile" / "inventory_profile.json"
        ledger_path = FORWARD / "evidence_ledger.json"
        portfolio_path = FORWARD / "portfolio.json"
        process_map_path = FORWARD / "process_map.json"
        review_path = FORWARD / "semantic_review.json"
        assessment_path = FORWARD / "customer_assessment.md"
        outcome_path = FORWARD / "outcome_review.json"

        profile = load_json(profile_path)
        ledger = load_json(ledger_path)
        portfolio = load_json(portfolio_path)
        mapping = load_json(process_map_path)
        review = load_json(review_path)
        self.assertEqual(validate_portfolio(portfolio, ledger, profile=profile), [])
        self.assertEqual(validate_process_map(mapping, profile, portfolio), [])
        hashes = {
            "inventory_profile_sha256": artifact_sha256(profile_path),
            "evidence_ledger_sha256": artifact_sha256(ledger_path),
            "portfolio_sha256": artifact_sha256(portfolio_path),
            "process_map_sha256": artifact_sha256(process_map_path),
        }
        self.assertEqual(
            validate_semantic_review(
                review,
                ledger,
                portfolio,
                mapping,
                profile,
                expected_hashes=hashes,
                today=date(2026, 7, 21),
                required_readiness="workshop_ready",
            ),
            [],
        )
        self.assertEqual(review["overall_readiness"], "workshop_ready")

        from render_customer_assessment import render

        rendered = render(profile, ledger, portfolio, mapping, review)
        expected = assessment_path.read_text(encoding="utf-8")
        self.assertEqual(rendered, expected)
        self.assertEqual(validate_customer_assessment(expected), [])
        self.assertLessEqual(markdown_word_count(expected), 900)
        self.assertIn("| Rank | Process | Why this order |", expected)
        self.assertIn("Latest record: 2026-07-02", expected)
        self.assertIn("Deployed: 4; Pipeline: 3; Paused: 1", expected)
        self.assertIn("5 analyst-mapped groups; customer confirmation required", expected)
        self.assertIn("| Unmapped | 3 records |", expected)
        self.assertLess(
            expected.index("### 1. Quarterly privileged access assurance"),
            expected.index("### 2. Supplier invoice exception resolution"),
        )
        self.assertLess(
            expected.index("### 2. Supplier invoice exception resolution"),
            expected.index("### 3. Vendor onboarding evidence readiness"),
        )
        self.assertIn("26,400 lines and 1,200 exceptions", expected)
        self.assertIn("18,000 submissions and 4,200 exceptions", expected)
        self.assertIn("Ground-truth owner: Access Governance Manager", expected)
        self.assertIn("cases matching final disposition / cases", expected)
        self.assertIn("Workshop ask:", expected)
        self.assertIn("thresholds from baselines and tolerances", expected)
        self.assertIn("No deployment or investment approval", expected)
        self.assertIn("Deferred pending owners, boundaries, or restart", expected)
        self.assertIn("missing documents flagged / known missing documents", expected)
        self.assertIn("Robot output: checklist and document gaps", expected)
        self.assertIn("Ground-truth owner: Supplier Risk Lead", expected)
        self.assertNotIn("Procurement Control Owner", expected)
        self.assertIn("Target: 2026-08-04", expected)
        self.assertIn("Decision: 2026-09-08", expected)
        self.assertNotIn("Agent ", expected)

        missing_linkage = copy.deepcopy(mapping)
        invoice = next(
            item
            for item in missing_linkage["orchestrations"]
            if item["opportunity_id"] == "OPP-INVOICE-EXCEPTION"
        )
        invoice_process = next(
            item
            for item in missing_linkage["processes"]
            if item["process_id"] == invoice["process_id"]
        )
        invoice_process["linkage"]["status"] = "validation_required"
        invoice["next_step"]["action"] = (
            "Confirm the boundary, assign the worklist owner, and select completed exceptions"
        )
        self.assertTrue(
            any(
                "must include linkage or identifier confirmation" in item
                for item in validate_process_map(missing_linkage, profile, portfolio)
            )
        )

        outcome = load_json(outcome_path)
        self.assertEqual(outcome["schema_version"], "1.0")
        self.assertTrue(outcome["review_id"].startswith("OUTCOME-REVIEW-"))
        self.assertEqual(outcome["assessment_sha256"], artifact_sha256(assessment_path))
        self.assertEqual(outcome["reviewed_at"], "2026-07-21")
        self.assertEqual(outcome["reviewer"]["mode"], "independent_agent")
        self.assertTrue(outcome["answer_blind"])
        self.assertEqual(
            set(outcome["scores"]),
            {
                "clarity",
                "process_specificity",
                "decision_utility",
                "account_team_actionability",
            },
        )
        for dimension, result in outcome["scores"].items():
            with self.subTest(dimension=dimension):
                score = result["score"]
                self.assertIsInstance(score, (int, float))
                self.assertNotIsInstance(score, bool)
                self.assertGreaterEqual(score, 4)
                self.assertLessEqual(score, 5)
                self.assertGreaterEqual(len(result["rationale"].split()), 5)

    @unittest.skipUnless(
        importlib.util.find_spec("docx"),
        "python-docx is required for draft DOCX verification",
    )
    def test_builder_fails_closed_without_page_renderer_and_marks_explicit_draft(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            case = create_case(root)
            common = [
                sys.executable,
                str(BUILDER),
                "--inventory-profile",
                str(case["profile_path"]),
                "--evidence-ledger",
                str(case["ledger_path"]),
                "--portfolio",
                str(case["portfolio_path"]),
                "--process-map",
                str(case["process_map_path"]),
                "--semantic-review",
                str(case["review_path"]),
                "--validation-date",
                case["date"].isoformat(),
                "--soffice",
                str(root / "missing-soffice"),
            ]
            ready = subprocess.run(
                [*common, "--output", str(root / "outputs" / "assessment.docx")],
                capture_output=True,
                text=True,
                check=False,
            )
            self.assertEqual(ready.returncode, 1)
            self.assertIn("soffice is required", ready.stderr)

            oversized = subprocess.run(
                [
                    *common,
                    "--output",
                    str(root / "outputs" / "oversized.docx"),
                    "--max-words",
                    "901",
                ],
                capture_output=True,
                text=True,
                check=False,
            )
            self.assertEqual(oversized.returncode, 1)
            self.assertIn("--max-words must be from 1 to 900", oversized.stderr)

            too_many_pages = subprocess.run(
                [
                    *common,
                    "--output",
                    str(root / "outputs" / "too-many-pages.docx"),
                    "--max-pages",
                    "3",
                ],
                capture_output=True,
                text=True,
                check=False,
            )
            self.assertEqual(too_many_pages.returncode, 1)
            self.assertIn("--max-pages must be 1 or 2", too_many_pages.stderr)

            protected_before = case["portfolio_path"].read_bytes()
            colliding_receipt = subprocess.run(
                [
                    *common,
                    "--output",
                    str(root / "draft-collision.docx"),
                    "--receipt",
                    str(case["portfolio_path"]),
                    "--draft-without-page-check",
                    "--force",
                ],
                capture_output=True,
                text=True,
                check=False,
            )
            self.assertEqual(colliding_receipt.returncode, 1)
            self.assertIn("collides with a protected input", colliding_receipt.stderr)
            self.assertEqual(case["portfolio_path"].read_bytes(), protected_before)

            draft_output = root / "draft.docx"
            draft = subprocess.run(
                [
                    *common,
                    "--output",
                    str(draft_output),
                    "--draft-without-page-check",
                ],
                capture_output=True,
                text=True,
                check=False,
            )
            self.assertEqual(draft.returncode, 0, draft.stderr)
            receipt = load_json(draft_output.with_suffix(".validation.json"))
            self.assertEqual(receipt["readiness"], "exploratory")
            self.assertFalse(receipt["layout_verified"])
            self.assertIsNone(receipt["page_count"])
            self.assertIsNone(receipt["rendered_pdf_sha256"])
            self.assertIsNone(receipt["rendered_pdf"])
            document = Document(draft_output)
            self.assertTrue(document.paragraphs[0].text.startswith("DRAFT - "))
            markdown_title = "Automation portfolio assessment: Lakeview Finance Authority"
            self.assertEqual(
                sum(paragraph.text == markdown_title for paragraph in document.paragraphs),
                0,
            )

    @unittest.skipUnless(
        importlib.util.find_spec("docx")
        and importlib.util.find_spec("pypdf")
        and shutil.which("soffice"),
        "python-docx, pypdf, and soffice are required for real page verification",
    )
    def test_three_recommendation_docx_is_branded_and_at_most_two_pages(self):
        from docx import Document
        from docx.oxml.ns import qn

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            case = create_case(root, opportunity_count=3)
            outputs = root / "outputs"
            output = outputs / "portfolio-assessment.docx"
            supporting_source = root / "strategy context.md"
            supporting_source.write_text(
                "Synthetic strategy context for receipt lineage.\n", encoding="utf-8"
            )
            result = subprocess.run(
                [
                    sys.executable,
                    str(BUILDER),
                    "--inventory-profile",
                    str(case["profile_path"]),
                    "--evidence-ledger",
                    str(case["ledger_path"]),
                    "--portfolio",
                    str(case["portfolio_path"]),
                    "--process-map",
                    str(case["process_map_path"]),
                    "--semantic-review",
                    str(case["review_path"]),
                    "--supporting-source",
                    str(supporting_source),
                    "--output",
                    str(output),
                    "--validation-date",
                    case["date"].isoformat(),
                    "--soffice",
                    shutil.which("soffice"),
                ],
                capture_output=True,
                text=True,
                check=False,
            )
            self.assertEqual(result.returncode, 0, result.stderr)
            receipt = load_json(output.with_suffix(".validation.json"))
            self.assertTrue(receipt["brand_verified"])
            self.assertTrue(receipt["layout_verified"])
            self.assertTrue(receipt["contracts_verified"])
            self.assertTrue(receipt["semantic_review_verified"])
            self.assertTrue(receipt["plain_language_verified"])
            self.assertEqual(
                receipt["input_hashes"]["semantic_review_sha256"],
                artifact_sha256(case["review_path"]),
            )
            self.assertIn(receipt["page_count"], (1, 2))
            self.assertRegex(receipt["rendered_pdf_sha256"], r"^[0-9a-f]{64}$")
            published_pdf = output.with_suffix(".pdf")
            self.assertTrue(published_pdf.exists())
            self.assertEqual(receipt["rendered_pdf"], published_pdf.name)
            self.assertEqual(
                receipt["rendered_pdf_sha256"], artifact_sha256(published_pdf)
            )
            if receipt["page_count"] == 2:
                from pypdf import PdfReader

                pages = PdfReader(str(published_pdf)).pages
                first_page_text = pages[0].extract_text()
                second_page_text = pages[1].extract_text()
                first_opportunity_id = case["portfolio"]["rankings"]["high_impact"][0]
                first_opportunity = next(
                    item
                    for item in case["portfolio"]["opportunities"]
                    if item["opportunity_id"] == first_opportunity_id
                )
                first_heading = f"1. {first_opportunity['name']}"
                self.assertIn("Top 3 Recommendations", first_page_text)
                self.assertNotIn(first_heading, first_page_text)
                self.assertIn(first_heading, second_page_text)
                second_opportunity_id = case["portfolio"]["rankings"]["high_impact"][1]
                second_opportunity = next(
                    item
                    for item in case["portfolio"]["opportunities"]
                    if item["opportunity_id"] == second_opportunity_id
                )
                second_heading = f"2. {second_opportunity['name']}"
                self.assertNotIn(second_heading, first_page_text)
                self.assertIn(second_heading, second_page_text)
                self.assertIn(
                    "approved vendor submission coordination pilot charter",
                    re.sub(r"\s+", " ", second_page_text),
                )
            self.assertEqual(receipt["recommendation_count"], 3)
            self.assertEqual(receipt["readiness"], "workshop_ready")
            document = Document(output)
            self.assertEqual(
                document.paragraphs[0].text,
                case["portfolio"]["customer_name"],
            )
            self.assertEqual(document.styles["Normal"].font.size.pt, 10)
            self.assertIn("Page", document.sections[0].footer.paragraphs[0].text)
            grid_columns = document.tables[0]._tbl.tblGrid.gridCol_lst
            widths = [int(column.get(qn("w:w"))) for column in grid_columns]
            self.assertEqual(len(widths), 2)
            self.assertGreater(widths[1], widths[0] * 2)
            self.assertEqual(
                receipt["ledger_inventory_as_of_date"],
                case["ledger"]["inventory_profile"]["as_of_date"],
            )
            self.assertEqual(
                receipt["latest_source_record_date"],
                case["profile"]["metadata"]["source_date_summary"]["latest_date"],
            )
            self.assertEqual(
                receipt["portfolio_as_of_date"], case["portfolio"]["as_of_date"]
            )
            self.assertEqual(
                receipt["raw_sources"],
                [
                    {
                        "kind": "inventory",
                        "name": case["profile"]["metadata"]["source_name"],
                        "sha256": case["profile"]["metadata"]["source_sha256"],
                    },
                    {
                        "kind": "supporting_context",
                        "name": supporting_source.name,
                        "sha256": artifact_sha256(supporting_source),
                    },
                ],
            )
            self.assertEqual(len(receipt["recommendation_evidence"]), 3)
            self.assertNotIn(str(root), json.dumps(receipt))


if __name__ == "__main__":
    unittest.main()
