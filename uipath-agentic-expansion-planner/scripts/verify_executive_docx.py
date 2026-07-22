#!/usr/bin/env python3
"""Verify the final UiPath agentic expansion executive DOCX artifact."""

from __future__ import annotations

import argparse
import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT


REQUIRED_HEADINGS = [
    "Executive Summary",
    "Source and Assumption Note",
    "Current Automation Footprint",
    "Public Strategy Alignment",
    "Prioritized Portfolio",
    "Top 5 High-Impact Recommendations",
    "Top 3 Low-Friction POC Candidates",
    "Value Framing",
    "Deployment and Governance Considerations",
    "Facts, Assumptions, and Validation Questions",
    "Workshop Prep",
    "Recommended Next Steps",
]
CUSTOMER_REQUIRED_HEADINGS = [
    "Source File Summary",
    "Current Automation Footprint",
    "Top 3 Recommendations",
]

REQUIRED_BRAND_COLORS = {"FA4616", "182126", "0BA2B3"}
FORBIDDEN_BRAND_COLORS = {"1F4E79"}
BRAND_FONT = "Arial"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Verify a rendered agentic expansion executive DOCX brief."
    )
    parser.add_argument("docx", type=Path, help="Rendered DOCX file to inspect")
    parser.add_argument(
        "--allow-landscape",
        action="store_true",
        help="Allow landscape orientation. By default, portrait is required.",
    )
    parser.add_argument(
        "--require-output-dir",
        action="store_true",
        help="Require the DOCX to be inside a directory named outputs.",
    )
    parser.add_argument(
        "--min-proposal-headings",
        type=int,
        default=5,
        help="Minimum level-3 proposal-card headings required after the Top 5 section.",
    )
    parser.add_argument(
        "--min-poc-headings",
        type=int,
        default=3,
        help="Minimum level-3 POC headings required after the Top 3 section.",
    )
    parser.add_argument(
        "--require-brand-style",
        action="store_true",
        help="Require UiPath-derived colors and shared-document font styling.",
    )
    parser.add_argument(
        "--profile",
        choices=("detailed", "customer-assessment"),
        default="detailed",
        help="Structural contract to verify",
    )
    parser.add_argument(
        "--rendered-pdf",
        type=Path,
        help="PDF rendered from this DOCX for page-count verification",
    )
    parser.add_argument("--max-pages", type=int, default=2)
    parser.add_argument(
        "--require-page-count",
        action="store_true",
        help="Fail unless a rendered PDF is supplied and within --max-pages",
    )
    return parser.parse_args()


def section_count(headings: list[tuple[int, str]], start: str, stop_candidates: set[str]) -> int:
    in_section = False
    count = 0
    for level, text in headings:
        if level == 2 and text == start:
            in_section = True
            continue
        if in_section and level == 2 and text in stop_candidates:
            break
        if in_section and level == 3:
            count += 1
    return count


def table_has_rank_header(document: Document) -> bool:
    for table in document.tables:
        if not table.rows:
            continue
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        if len(headers) >= 2 and headers[0] == "rank" and headers[1] == "opportunity":
            return True
    return False


def docx_xml_text(path: Path) -> str:
    chunks: list[str] = []
    with zipfile.ZipFile(path) as archive:
        for name in archive.namelist():
            if name.startswith("word/") and name.endswith(".xml"):
                chunks.append(archive.read(name).decode("utf-8", errors="ignore"))
    return "\n".join(chunks).upper()


def brand_style_failures(path: Path) -> list[str]:
    xml = docx_xml_text(path)
    failures: list[str] = []
    missing = sorted(color for color in REQUIRED_BRAND_COLORS if color not in xml)
    if missing:
        failures.append("Missing required UiPath brand color(s): " + ", ".join(missing))
    forbidden = sorted(color for color in FORBIDDEN_BRAND_COLORS if color in xml)
    if forbidden:
        failures.append("Old generic Office-style color(s) still present: " + ", ".join(forbidden))
    if BRAND_FONT.upper() not in xml:
        failures.append(f"Missing shared-document fallback font: {BRAND_FONT}")
    if "APTOS" in xml:
        failures.append("Aptos font remains in generated DOCX styling.")
    return failures


def pdf_page_count(path: Path) -> int:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency is pinned for repository validation
        raise RuntimeError("pypdf is required for rendered page verification") from exc
    return len(PdfReader(str(path)).pages)


def main() -> int:
    args = parse_args()
    failures: list[str] = []

    if args.profile == "customer-assessment" and not 1 <= args.max_pages <= 2:
        failures.append("Customer-assessment --max-pages must be 1 or 2.")

    if not args.docx.exists():
        failures.append(f"DOCX does not exist: {args.docx}")
        print("\n".join(failures), file=sys.stderr)
        return 1

    if args.require_output_dir and args.docx.parent.name != "outputs":
        failures.append(f"DOCX must be in an outputs directory: {args.docx}")

    document = Document(args.docx)
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    headings: list[tuple[int, str]] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if style_name.startswith("Heading"):
            try:
                level = int(style_name.split()[-1])
            except ValueError:
                level = 0
            headings.append((level, text))

    if not paragraphs:
        failures.append("DOCX has no non-empty paragraphs.")
    else:
        title = paragraphs[0]
        if len(title) < 8:
            failures.append(f"DOCX title looks empty or too short: {title!r}")

    if not args.allow_landscape and document.sections[0].orientation != WD_ORIENT.PORTRAIT:
        failures.append("DOCX must be portrait unless landscape was explicitly requested.")

    heading_texts = [text for _, text in headings]
    proposal_count = 0
    poc_count = 0
    if args.profile == "customer-assessment":
        level_two = [text for level, text in headings if level == 2]
        if level_two != CUSTOMER_REQUIRED_HEADINGS:
            failures.append(
                "Customer assessment must contain exactly these headings in order: "
                + ", ".join(CUSTOMER_REQUIRED_HEADINGS)
            )
        proposal_count = section_count(headings, "Top 3 Recommendations", set())
        if not 1 <= proposal_count <= 3:
            failures.append(
                f"Customer assessment must contain one to three recommendations; found {proposal_count}."
            )
        recommendation_headings = [text for level, text in headings if level == 3]
        for index, heading in enumerate(recommendation_headings, start=1):
            if not re.match(rf"^{index}\.\s+\S", heading):
                failures.append(
                    "Customer recommendation headings must be numbered consecutively from 1."
                )
                break
        if len(document.tables) < 1:
            failures.append("Customer assessment must contain the footprint table.")
        all_text = "\n".join(paragraphs)
        all_text += "\n" + "\n".join(
            cell.text for table in document.tables for row in table.rows for cell in row.cells
        )
        if re.search(
            r"\b(?:INV|SRC|ASM|OPP|PROC|LEDGER|PORTFOLIO|PROCESS-MAP|REVIEW)-[A-Z0-9]",
            all_text,
            re.I,
        ):
            failures.append("Customer assessment exposes an internal evidence or contract ID.")
        if re.search(
            r"(?:^|[\s(])(?:/(?:Users|home|private|tmp|var)/[^\s)]+|[A-Za-z]:\\[^\s)]+)",
            all_text,
            re.I,
        ):
            failures.append("Customer assessment exposes a local filesystem path.")
        if re.search(r"\b[0-9a-f]{64}\b", all_text, re.I):
            failures.append("Customer assessment exposes a raw SHA-256 value.")
    else:
        missing = [heading for heading in REQUIRED_HEADINGS if heading not in heading_texts]
        if missing:
            failures.append("Missing required headings: " + ", ".join(missing))
        if not any("Source Ledger" in text or text.startswith("Appendix") for text in heading_texts):
            failures.append("Missing appendix/source-ledger heading.")
        if len(document.tables) < 3:
            failures.append(f"Expected at least 3 tables; found {len(document.tables)}.")
        if not table_has_rank_header(document):
            failures.append("No prioritized portfolio table with Rank / Opportunity header found.")
        proposal_count = section_count(
            headings,
            "Top 5 High-Impact Recommendations",
            {"Top 3 Low-Friction POC Candidates"},
        )
        if proposal_count < args.min_proposal_headings:
            failures.append(
                f"Expected at least {args.min_proposal_headings} proposal headings; found {proposal_count}."
            )
        poc_count = section_count(
            headings,
            "Top 3 Low-Friction POC Candidates",
            {"Value Framing"},
        )
        if poc_count < args.min_poc_headings:
            failures.append(
                f"Expected at least {args.min_poc_headings} POC headings; found {poc_count}."
            )
        first_heading_names = [text for _, text in headings[:6]]
        if any("Source Ledger" in text or text.startswith("Appendix") for text in first_heading_names):
            failures.append("Source ledger appears too early; keep it in the appendix.")

    page_count = None
    if args.rendered_pdf:
        if not args.rendered_pdf.exists():
            failures.append(f"Rendered PDF does not exist: {args.rendered_pdf}")
        else:
            try:
                page_count = pdf_page_count(args.rendered_pdf)
            except RuntimeError as exc:
                failures.append(str(exc))
            if page_count is not None and not 1 <= page_count <= args.max_pages:
                failures.append(
                    f"Rendered customer assessment has {page_count} page(s); maximum is {args.max_pages}."
                )
    elif args.require_page_count:
        failures.append("--require-page-count requires --rendered-pdf")

    if args.require_brand_style:
        failures.extend(brand_style_failures(args.docx))

    if failures:
        for failure in failures:
            print(f"FAIL: {failure}", file=sys.stderr)
        return 1

    orientation = (
        "PORTRAIT" if document.sections[0].orientation == WD_ORIENT.PORTRAIT else "LANDSCAPE"
    )
    print(f"OK: {args.docx}")
    print(f"title={paragraphs[0] if paragraphs else ''}")
    print(f"orientation={orientation}")
    print(f"paragraphs={len(document.paragraphs)}")
    print(f"tables={len(document.tables)}")
    print(f"proposal_headings={proposal_count}")
    print(f"poc_headings={poc_count}")
    if page_count is not None:
        print(f"pages={page_count}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
