#!/usr/bin/env python3
"""Render a UiPath agentic expansion Markdown brief as a polished executive Word document."""

from __future__ import annotations

import argparse
import re
from datetime import date
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from validate_customer_assessment import markdown_word_count


UIPATH_ORANGE = RGBColor(250, 70, 22)
UIPATH_DEEP_BLUE = RGBColor(24, 33, 38)
UIPATH_AGENTIC_TEAL = RGBColor(11, 162, 179)
UIPATH_GRAY_TEXT = RGBColor(72, 72, 72)
UIPATH_ORANGE_HEX = "FA4616"
UIPATH_DEEP_BLUE_HEX = "182126"
UIPATH_AGENTIC_TEAL_HEX = "0BA2B3"
UIPATH_ROW_ALT = "F6F6F6"
BRAND_FONT = "Arial"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Convert a concise UiPath agentic expansion Markdown executive brief to .docx."
    )
    parser.add_argument("markdown", type=Path, help="Input Markdown briefing file")
    parser.add_argument("docx", type=Path, help="Output .docx file")
    parser.add_argument("--title", help="Override the document title")
    parser.add_argument(
        "--subtitle",
        default=f"Executive briefing | Prepared {date.today().isoformat()}",
        help="Subtitle shown under the title",
    )
    parser.add_argument(
        "--portrait",
        action="store_true",
        help="Keep portrait orientation. Accepted for backward compatibility; portrait is the default.",
    )
    parser.add_argument(
        "--auto-landscape",
        action="store_true",
        help="Opt in to landscape orientation when wide tables are present. Ignored when --portrait is also set.",
    )
    parser.add_argument(
        "--max-words",
        type=int,
        default=None,
        help="Word ceiling; defaults to 900 for customer-assessment and 3200 otherwise",
    )
    parser.add_argument(
        "--profile",
        choices=("detailed", "customer-assessment"),
        default="detailed",
        help="Document density and footer profile",
    )
    return parser.parse_args()


def split_table_row(line: str) -> list[str]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [cell.strip() for cell in stripped.split("|")]


def is_separator(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def is_table_start(lines: Sequence[str], index: int) -> bool:
    return (
        index + 1 < len(lines)
        and "|" in lines[index]
        and lines[index].strip().startswith("|")
        and is_separator(lines[index + 1])
    )


def is_heading(line: str) -> bool:
    return bool(re.match(r"^#{1,4}\s+", line.strip()))


def is_bullet(line: str) -> bool:
    return bool(re.match(r"^\s*[-*]\s+", line))


def is_numbered(line: str) -> bool:
    return bool(re.match(r"^\s*\d+\.\s+", line))


def iter_blocks(lines: Sequence[str]):
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            index += 1
            continue

        if is_table_start(lines, index):
            rows = [split_table_row(lines[index])]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append(split_table_row(lines[index]))
                index += 1
            yield ("table", rows)
            continue

        heading_match = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if heading_match:
            yield ("heading", (len(heading_match.group(1)), heading_match.group(2).strip()))
            index += 1
            continue

        if is_bullet(line):
            items = []
            while index < len(lines) and is_bullet(lines[index]):
                items.append(re.sub(r"^\s*[-*]\s+", "", lines[index].strip()))
                index += 1
            yield ("bullets", items)
            continue

        if is_numbered(line):
            items = []
            while index < len(lines) and is_numbered(lines[index]):
                items.append(re.sub(r"^\s*\d+\.\s+", "", lines[index].strip()))
                index += 1
            yield ("numbers", items)
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            next_line = lines[index].rstrip()
            if (
                not next_line.strip()
                or is_heading(next_line)
                or is_table_start(lines, index)
                or is_bullet(next_line)
                or is_numbered(next_line)
            ):
                break
            paragraph_lines.append(next_line.strip())
            index += 1
        yield ("paragraph", " ".join(paragraph_lines))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 80, bottom: int = 80, end: int = 80) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        element = tc_mar.find(qn(f"w:{margin_name}"))
        if element is None:
            element = OxmlElement(f"w:{margin_name}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(margin_value))
        element.set(qn("w:type"), "dxa")


def set_font(run, *, name: str = BRAND_FONT) -> None:
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("ascii", "hAnsi", "cs"):
        r_fonts.set(qn(f"w:{attr}"), name)


def add_inline_markdown(paragraph, text: str, *, font_size: float | None = None) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.color.rgb = RGBColor(75, 75, 75)
        else:
            run = paragraph.add_run(part)
        if not (part.startswith("`") and part.endswith("`")):
            set_font(run)
        if font_size is not None:
            run.font.size = Pt(font_size)


def style_document(document: Document, *, landscape: bool, profile: str = "detailed") -> None:
    section = document.sections[0]
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    compact = profile == "customer-assessment"
    section.top_margin = Inches(0.45 if compact else 0.6)
    section.bottom_margin = Inches(0.42 if compact else 0.55)
    section.left_margin = Inches(0.55 if compact or landscape else 0.7)
    section.right_margin = Inches(0.55 if compact or landscape else 0.7)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = BRAND_FONT
    normal.font.size = Pt(10 if compact else 9.5)

    for style_name, size, color in [
        ("Heading 1", 14 if compact else 16, UIPATH_ORANGE),
        ("Heading 2", 10.5 if compact else 12, UIPATH_DEEP_BLUE),
        ("Heading 3", 9.5 if compact else 10.5, UIPATH_AGENTIC_TEAL),
    ]:
        style = styles[style_name]
        style.font.name = BRAND_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(6 if compact else 10)
        style.paragraph_format.space_after = Pt(2 if compact else 4)

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER if compact else WD_ALIGN_PARAGRAPH.RIGHT
    )
    label = footer.add_run(
        "Page "
        if compact
        else "Agentic expansion briefing | Page "
    )
    label.font.size = Pt(8)
    label.font.color.rgb = UIPATH_GRAY_TEXT
    set_font(label)
    begin_run = footer.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin)
    instruction_run = footer.add_run()
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    instruction_run._r.append(instruction)
    separate_run = footer.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)
    result = footer.add_run("1")
    result.font.size = Pt(8)
    result.font.color.rgb = UIPATH_GRAY_TEXT
    set_font(result)
    end_run = footer.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def add_title_block(
    document: Document, title: str, subtitle: str, *, profile: str = "detailed"
) -> None:
    compact = profile == "customer-assessment"
    title_para = document.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(17 if compact else 20)
    title_run.font.color.rgb = UIPATH_ORANGE
    set_font(title_run)

    subtitle_para = document.add_paragraph()
    subtitle_run = subtitle_para.add_run(subtitle)
    subtitle_run.font.size = Pt(8.5 if compact else 9.5)
    subtitle_run.font.color.rgb = UIPATH_GRAY_TEXT
    set_font(subtitle_run)

    rule = document.add_paragraph()
    rule.paragraph_format.space_after = Pt(4 if compact else 8)
    run = rule.add_run(" ")
    run.font.size = Pt(1)
    p_pr = rule._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), UIPATH_DEEP_BLUE_HEX)
    border.append(bottom)
    p_pr.append(border)


def add_table(
    document: Document, rows: Sequence[Sequence[str]], *, profile: str = "detailed"
) -> None:
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = not (
        profile == "customer-assessment" and column_count in {2, 3}
    )
    if profile == "customer-assessment" and column_count in {2, 3}:
        section = document.sections[0]
        usable_width = section.page_width - section.left_margin - section.right_margin
        ratios = (0.24, 0.76) if column_count == 2 else (0.08, 0.28, 0.64)
        for column_index, ratio in enumerate(ratios):
            column_width = int(usable_width * ratio)
            table.columns[column_index].width = column_width
            for cell in table.columns[column_index].cells:
                cell.width = column_width

    for row_index, row in enumerate(rows):
        for col_index in range(column_count):
            text = row[col_index] if col_index < len(row) else ""
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell, top=45, start=55, bottom=45, end=55) if profile == "customer-assessment" else set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, UIPATH_DEEP_BLUE_HEX)
            elif row_index % 2 == 0:
                set_cell_shading(cell, UIPATH_ROW_ALT)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            font_size = (
                (7.8 if column_count == 3 else 8.3)
                if profile == "customer-assessment"
                else 7.8 if column_count > 5 else 8.5
            )
            add_inline_markdown(paragraph, text, font_size=font_size)
            for run in paragraph.runs:
                if row_index == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                elif profile == "customer-assessment" and col_index == 0:
                    run.bold = True
                    run.font.color.rgb = UIPATH_DEEP_BLUE

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0 if profile == "customer-assessment" else 4)


def render_blocks(
    document: Document,
    blocks: Iterable[tuple[str, object]],
    title: str,
    *,
    profile: str = "detailed",
) -> None:
    compact = profile == "customer-assessment"
    first_h1_skipped = False
    block_list = list(blocks)
    recommendation_count = sum(
        1
        for block_type, value in block_list
        if block_type == "heading" and value[0] == 3
    )
    recommendation_index = 0
    for block_type, value in block_list:
        if block_type == "heading":
            level, text = value
            if (
                level == 1
                and not first_h1_skipped
                and (compact or text.strip() == title.strip())
            ):
                first_h1_skipped = True
                continue
            style = "Heading 1" if level <= 1 else "Heading 2" if level == 2 else "Heading 3"
            paragraph = document.add_paragraph(text, style=style)
            if compact and level == 2 and text.strip() == "Top 3 Recommendations":
                paragraph.paragraph_format.keep_with_next = True
            if compact and level == 3:
                recommendation_index += 1
                paragraph.paragraph_format.keep_with_next = True
                if (
                    recommendation_count == 3
                    and recommendation_index == 1
                ) or (
                    recommendation_count == 2
                    and recommendation_index == 2
                ):
                    paragraph.paragraph_format.page_break_before = True
        elif block_type == "paragraph":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(2 if compact else 5)
            paragraph.paragraph_format.line_spacing = 1.0 if compact else 1.05
            add_inline_markdown(paragraph, value)
        elif block_type == "bullets":
            for item in value:
                paragraph = document.add_paragraph(style="List Bullet")
                paragraph.paragraph_format.space_after = Pt(2 if compact else 2)
                if compact:
                    paragraph.paragraph_format.left_indent = Inches(0.16)
                    paragraph.paragraph_format.first_line_indent = Inches(-0.11)
                    paragraph.paragraph_format.line_spacing = 1.0
                add_inline_markdown(paragraph, item)
                if compact and paragraph.runs and paragraph.runs[0].bold:
                    paragraph.runs[0].font.color.rgb = UIPATH_DEEP_BLUE
        elif block_type == "numbers":
            for item in value:
                paragraph = document.add_paragraph(style="List Number")
                paragraph.paragraph_format.space_after = Pt(1 if compact else 2)
                add_inline_markdown(paragraph, item)
        elif block_type == "table":
            add_table(document, value, profile=profile)


def first_heading_one(blocks: Sequence[tuple[str, object]]) -> str | None:
    for block_type, value in blocks:
        if block_type == "heading":
            level, text = value
            if level == 1:
                return text
    return None


def has_wide_table(blocks: Sequence[tuple[str, object]]) -> bool:
    for block_type, value in blocks:
        if block_type == "table" and value and max(len(row) for row in value) > 5:
            return True
    return False


def main() -> int:
    args = parse_args()
    markdown_text = args.markdown.read_text(encoding="utf-8")
    word_count = markdown_word_count(markdown_text)
    max_words = args.max_words or (900 if args.profile == "customer-assessment" else 3200)
    if args.profile == "customer-assessment" and max_words > 900:
        print("FAIL: customer-assessment --max-words cannot exceed 900.")
        return 1
    if word_count > max_words and args.profile == "customer-assessment":
        print(
            f"FAIL: {args.markdown} has {word_count} words; customer-assessment maximum is "
            f"{max_words}."
        )
        return 1
    if word_count > max_words:
        print(
            f"Warning: {args.markdown} has {word_count} words. "
            "For executive .docx output, consider shortening before sharing."
        )

    blocks = list(iter_blocks(markdown_text.splitlines()))
    title = args.title or first_heading_one(blocks) or args.markdown.stem.replace("_", " ").title()

    document = Document()
    use_landscape = has_wide_table(blocks) and args.auto_landscape and not args.portrait
    style_document(document, landscape=use_landscape, profile=args.profile)
    add_title_block(document, title, args.subtitle, profile=args.profile)
    render_blocks(document, blocks, title, profile=args.profile)

    args.docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.docx)
    print(f"Wrote {args.docx}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
