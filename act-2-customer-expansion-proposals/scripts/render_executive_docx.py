#!/usr/bin/env python3
"""Render an Act 2 Markdown brief as a polished executive Word document."""

from __future__ import annotations

import argparse
import re
from datetime import date
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ACCENT = RGBColor(31, 78, 121)
ACCENT_DARK = "1F4E79"
ACCENT_LIGHT = "D9EAF7"
GRAY_TEXT = RGBColor(82, 82, 82)
ROW_ALT = "F7F9FB"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Convert a concise Act 2 Markdown executive brief to .docx."
    )
    parser.add_argument("markdown", type=Path, help="Input Markdown briefing file")
    parser.add_argument("docx", type=Path, help="Output .docx file")
    parser.add_argument("--title", help="Override the document title")
    parser.add_argument(
        "--subtitle",
        default=f"Executive briefing | Prepared {date.today().isoformat()}",
        help="Subtitle shown under the title",
    )
    parser.add_argument(
        "--portrait",
        action="store_true",
        help="Keep portrait orientation. Accepted for backward compatibility; portrait is the default.",
    )
    parser.add_argument(
        "--auto-landscape",
        action="store_true",
        help="Opt in to landscape orientation when wide tables are present. Ignored when --portrait is also set.",
    )
    parser.add_argument(
        "--max-words",
        type=int,
        default=3200,
        help="Warn when the Markdown exceeds this word count",
    )
    return parser.parse_args()


def split_table_row(line: str) -> list[str]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [cell.strip() for cell in stripped.split("|")]


def is_separator(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def is_table_start(lines: Sequence[str], index: int) -> bool:
    return (
        index + 1 < len(lines)
        and "|" in lines[index]
        and lines[index].strip().startswith("|")
        and is_separator(lines[index + 1])
    )


def is_heading(line: str) -> bool:
    return bool(re.match(r"^#{1,4}\s+", line.strip()))


def is_bullet(line: str) -> bool:
    return bool(re.match(r"^\s*[-*]\s+", line))


def is_numbered(line: str) -> bool:
    return bool(re.match(r"^\s*\d+\.\s+", line))


def iter_blocks(lines: Sequence[str]):
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            index += 1
            continue

        if is_table_start(lines, index):
            rows = [split_table_row(lines[index])]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append(split_table_row(lines[index]))
                index += 1
            yield ("table", rows)
            continue

        heading_match = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if heading_match:
            yield ("heading", (len(heading_match.group(1)), heading_match.group(2).strip()))
            index += 1
            continue

        if is_bullet(line):
            items = []
            while index < len(lines) and is_bullet(lines[index]):
                items.append(re.sub(r"^\s*[-*]\s+", "", lines[index].strip()))
                index += 1
            yield ("bullets", items)
            continue

        if is_numbered(line):
            items = []
            while index < len(lines) and is_numbered(lines[index]):
                items.append(re.sub(r"^\s*\d+\.\s+", "", lines[index].strip()))
                index += 1
            yield ("numbers", items)
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            next_line = lines[index].rstrip()
            if (
                not next_line.strip()
                or is_heading(next_line)
                or is_table_start(lines, index)
                or is_bullet(next_line)
                or is_numbered(next_line)
            ):
                break
            paragraph_lines.append(next_line.strip())
            index += 1
        yield ("paragraph", " ".join(paragraph_lines))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 80, bottom: int = 80, end: int = 80) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        element = tc_mar.find(qn(f"w:{margin_name}"))
        if element is None:
            element = OxmlElement(f"w:{margin_name}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(margin_value))
        element.set(qn("w:type"), "dxa")


def add_inline_markdown(paragraph, text: str, *, font_size: float | None = None) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.color.rgb = RGBColor(75, 75, 75)
        else:
            run = paragraph.add_run(part)
        if font_size is not None:
            run.font.size = Pt(font_size)


def style_document(document: Document, *, landscape: bool) -> None:
    section = document.sections[0]
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55 if landscape else 0.7)
    section.right_margin = Inches(0.55 if landscape else 0.7)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)

    for style_name, size, color in [
        ("Heading 1", 16, ACCENT),
        ("Heading 2", 12, ACCENT),
        ("Heading 3", 10.5, ACCENT),
    ]:
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)

    footer = section.footer.paragraphs[0]
    footer.text = "Act 2 expansion briefing"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = GRAY_TEXT


def add_title_block(document: Document, title: str, subtitle: str) -> None:
    title_para = document.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = ACCENT

    subtitle_para = document.add_paragraph()
    subtitle_run = subtitle_para.add_run(subtitle)
    subtitle_run.font.size = Pt(9.5)
    subtitle_run.font.color.rgb = GRAY_TEXT

    rule = document.add_paragraph()
    rule.paragraph_format.space_after = Pt(8)
    run = rule.add_run(" ")
    run.font.size = Pt(1)
    p_pr = rule._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), ACCENT_DARK)
    border.append(bottom)
    p_pr.append(border)


def add_table(document: Document, rows: Sequence[Sequence[str]]) -> None:
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True

    for row_index, row in enumerate(rows):
        for col_index in range(column_count):
            text = row[col_index] if col_index < len(row) else ""
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, ACCENT_DARK)
            elif row_index % 2 == 0:
                set_cell_shading(cell, ROW_ALT)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline_markdown(paragraph, text, font_size=7.8 if column_count > 5 else 8.5)
            for run in paragraph.runs:
                if row_index == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)

    document.add_paragraph()


def render_blocks(document: Document, blocks: Iterable[tuple[str, object]], title: str) -> None:
    first_h1_skipped = False
    for block_type, value in blocks:
        if block_type == "heading":
            level, text = value
            if level == 1 and text.strip() == title.strip() and not first_h1_skipped:
                first_h1_skipped = True
                continue
            style = "Heading 1" if level <= 1 else "Heading 2" if level == 2 else "Heading 3"
            document.add_paragraph(text, style=style)
        elif block_type == "paragraph":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(5)
            paragraph.paragraph_format.line_spacing = 1.05
            add_inline_markdown(paragraph, value)
        elif block_type == "bullets":
            for item in value:
                paragraph = document.add_paragraph(style="List Bullet")
                paragraph.paragraph_format.space_after = Pt(2)
                add_inline_markdown(paragraph, item)
        elif block_type == "numbers":
            for item in value:
                paragraph = document.add_paragraph(style="List Number")
                paragraph.paragraph_format.space_after = Pt(2)
                add_inline_markdown(paragraph, item)
        elif block_type == "table":
            add_table(document, value)


def first_heading_one(blocks: Sequence[tuple[str, object]]) -> str | None:
    for block_type, value in blocks:
        if block_type == "heading":
            level, text = value
            if level == 1:
                return text
    return None


def has_wide_table(blocks: Sequence[tuple[str, object]]) -> bool:
    for block_type, value in blocks:
        if block_type == "table" and value and max(len(row) for row in value) > 5:
            return True
    return False


def main() -> int:
    args = parse_args()
    markdown_text = args.markdown.read_text(encoding="utf-8")
    word_count = len(re.findall(r"\b\w+\b", markdown_text))
    if word_count > args.max_words:
        print(
            f"Warning: {args.markdown} has {word_count} words. "
            "For executive .docx output, consider shortening before sharing."
        )

    blocks = list(iter_blocks(markdown_text.splitlines()))
    title = args.title or first_heading_one(blocks) or args.markdown.stem.replace("_", " ").title()

    document = Document()
    use_landscape = has_wide_table(blocks) and args.auto_landscape and not args.portrait
    style_document(document, landscape=use_landscape)
    add_title_block(document, title, args.subtitle)
    render_blocks(document, blocks, title)

    args.docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.docx)
    print(f"Wrote {args.docx}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
